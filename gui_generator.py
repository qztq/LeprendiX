import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import sqlite3
import datetime
from PIL import Image, ImageTk
from docx import Document
import os
import requests 
import subprocess 
import sys
import shutil
from docx.enum.text import WD_UNDERLINE
from docx.shared import Pt
import calendar # Am Anfang der Datei zu den anderen Imports hinzufügen
import logging
import threading
from config_loader import CONFIG
from geopy.geocoders import Nominatim
from geopy.exc import GeocoderTimedOut, GeocoderUnavailable
import mobile_connect



# --- HELPER FUNCTIONS FOR PATH RESOLUTION ---
def resolve_data_path(relative_path):
    """Resolves path for writable data (next to exe or source dir)."""
    if getattr(sys, 'frozen', False):
        base_path = os.path.dirname(sys.executable)
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)

def resolve_resource_path(relative_path):
    """Resolves path for read-only resources (bundled in exe or source dir)."""
    # 1. Try bundled path (PyInstaller)
    if hasattr(sys, '_MEIPASS'):
        bundled_path = os.path.join(sys._MEIPASS, relative_path)
        if os.path.exists(bundled_path):
            return bundled_path
    # 2. Fallback to data path
    return resolve_data_path(relative_path)

# --- KONFIGURATION ---
DATABASE_NAME = resolve_data_path('patienten.db')
TEMPLATE_FILE = resolve_resource_path('honorar_vorlage.docx') 


def start_gui():
    root = tk.Tk()
    app = HonorarGeneratorApp(root)
    root.mainloop()
    

def _ensure_status_column():
    """
    Stellt sicher, dass die Spalte 'invoiced_since_reset' in der patienten-Tabelle existiert.
    Wird einmal beim Start aufgerufen.
    """
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    # Check if table exists to avoid crash on first run
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='patienten'")
    if not cursor.fetchone():
        conn.close()
        return

    try:
        # Versuche, die Spalte zu lesen
        cursor.execute("SELECT invoiced_since_reset FROM patienten LIMIT 1")
    except sqlite3.OperationalError:
        # Wenn die Spalte nicht existiert, füge sie hinzu
        cursor.execute("ALTER TABLE patienten ADD COLUMN invoiced_since_reset INTEGER DEFAULT 0")
        conn.commit()
        logging.info("Spalte 'invoiced_since_reset' in patienten-Tabelle hinzugefügt.")

    try:
        cursor.execute("SELECT last_selected_kurznamen FROM patienten LIMIT 1")
    except sqlite3.OperationalError:
        cursor.execute("ALTER TABLE patienten ADD COLUMN last_selected_kurznamen TEXT DEFAULT ''")
        conn.commit()
        logging.info("Spalte 'last_selected_kurznamen' zur patienten-Tabelle hinzugefügt.")
    
    
    try:
        cursor.execute("SELECT is_archived FROM patienten LIMIT 1")
    except sqlite3.OperationalError:
        cursor.execute("ALTER TABLE patienten ADD COLUMN is_archived INTEGER DEFAULT 0")
        conn.commit()
        logging.info("Spalte 'is_archived' zur patienten-Tabelle hinzugefügt.")
    finally:
        conn.close()

def _update_invoiced_status(patient_id, status=1):
    """Setzt den Status eines Patienten auf 1 (Grün/Abgerechnet) oder 0 (Rot/Offen)."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    try:
        cursor.execute("UPDATE patienten SET invoiced_since_reset = ? WHERE id = ?", (status, patient_id))
        conn.commit()
        logging.info(f"Patient {patient_id} Honorarnoten-Status auf {status} gesetzt.")
    except Exception as e:
        logging.error(f"FEHLER beim Status-Update für Patient {patient_id}: {e}")
    finally:
        conn.close()
        
# Sicherstellen, dass die Spalte beim Start existiert
_ensure_status_column()

def print_document_silently(file_path):
    """
    Versucht, die angegebene Datei direkt an den Standarddrucker zu senden.
    """
    if not os.path.exists(file_path):
        return False, "Datei zum Drucken nicht gefunden."

    try:
        if sys.platform.startswith('win'):
            # Windows: Nutzt den 'print' Verb des Dateityps (oft dialoglos)
            os.startfile(file_path, 'print')
            return True, "Druckauftrag an Windows-Standarddrucker gesendet."
            
        elif sys.platform.startswith('darwin') or sys.platform.startswith('linux'):
            # macOS/Linux: Nutzt 'lpr' (kann auf einigen Systemen einen Dialog auslösen)
            subprocess.run(['lpr', file_path], check=True)
            return True, "Druckauftrag via lpr (Linux/macOS) gesendet."
            
        else:
            return False, f"Automatisches Drucken wird auf dem Betriebssystem '{sys.platform}' nicht unterstützt."
            
    except subprocess.CalledProcessError as e:
        return False, f"Fehler beim LPR-Druckbefehl: {e}"
    except Exception as e:
        return False, f"Fehler beim direkten Drucken: {e}"

def get_patient_data(search_name):
    """Sucht Patienten und gibt ID und alle Adressfelder (inkl. Kilometergeld und letzte Leistungen) zurück."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    search_term = f'%{search_name}%'
    # Wichtig: last_selected_kurznamen ist Spalte 13 (Index 12)
    # Wichtig: invoiced_since_reset ist Spalte 14 (Index 13), wird hier nicht benötigt, aber später
    
    query = """
    SELECT id, vorname, nachname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, last_selected_kurznamen
    FROM patienten 
    WHERE (nachname LIKE ? OR vorname LIKE ?) AND (is_archived IS NULL OR is_archived = 0)
    """
    params = [search_term, search_term]
    
    if search_name.isdigit():
        query += " OR (id = ? AND (is_archived IS NULL OR is_archived = 0))"
        params.append(search_name)
        
    cursor.execute(query, params)
    results = cursor.fetchall()
    conn.close()
    return results

def save_last_selected_leistungen(patient_id, kurznamen_set):
    """Speichert die Liste der zuletzt ausgewählten Kurznamen für einen Patienten in der DB."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    # Konvertiere das Set in einen String
    kurznamen_str = ','.join(sorted(list(kurznamen_set)))
    try:
        cursor.execute("""
        UPDATE patienten 
        SET last_selected_kurznamen = ?
        WHERE id = ?
        """, (kurznamen_str, patient_id))
        conn.commit()
    except Exception as e:
        logging.error(f"Fehler beim Speichern der letzten Leistungen für Patient {patient_id}: {e}")
    finally:
        conn.close()

def get_patient_leistungen(patient_id):
    """Holt alle NICHT ABGERECHNETEN Leistungen für die GUI-Anzeige."""
    # TODO: Muss später um 'WHERE abgerechnet_am IS NULL' erweitert werden
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute("""
    SELECT id, datum, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag 
    FROM leistungen 
    WHERE patient_id = ? 
    ORDER BY datum ASC, uhrzeit_von ASC
    """, (patient_id,))
    leistungen = cursor.fetchall()
    conn.close()
    return leistungen

def get_patient_leistungen_for_template(patient_id):
    """Holt NICHT ABGERECHNETE Leistungen (ohne ID) für die Word-Generierung."""
    # TODO: Muss später um 'WHERE abgerechnet_am IS NULL' erweitert werden
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    cursor.execute("""
    SELECT datum, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag 
    FROM leistungen 
    WHERE patient_id = ? 
    ORDER BY datum ASC, uhrzeit_von ASC
    """, (patient_id,))
    leistungen = cursor.fetchall()
    conn.close()
    return leistungen

def get_all_stammdaten_dict(archived=False):
    """Holt alle Stammdaten aus der DB und gibt sie als Liste und Dict zurück."""
    conn = sqlite3.connect(DATABASE_NAME)
    cursor = conn.cursor()
    
    # NEU: Abhängig vom Flag archivierte oder aktive laden
    if archived:
        query = "SELECT kurzname, beschreibung, standard_betrag FROM stammdaten_leistungen WHERE is_archived = 1 ORDER BY kurzname"
    else:
        query = "SELECT kurzname, beschreibung, standard_betrag FROM stammdaten_leistungen WHERE is_archived = 0 ORDER BY kurzname"
        
    cursor.execute(query)
    results = cursor.fetchall()
    conn.close()
    
    stammdaten_list = [f"{r[0]} - {r[1]}" for r in results]
    stammdaten_dict = {item: r[2] for item, r in zip(stammdaten_list, results)}
    return stammdaten_list, stammdaten_dict

def search_teamup_events(search_term, start_date=None, end_date=None, mode='standard'):
    """
    Sucht Teamup-Kalendereinträge basierend auf dem Titel/Notizen.
    """
    
    # Lade Konfiguration dynamisch
    api_key = CONFIG.get('TEAMUP_API_KEY', '')
    calendar_id = CONFIG.get('TEAMUP_CALENDAR_ID', '')
    
    clean_api_key = api_key.strip()
    base_url = f"https://api.teamup.com/{calendar_id}/events"
    
    headers = {
        'Teamup-Token': clean_api_key, 
        'Accept': 'application/json'
    }
    
    # BEGRENZUNG DES ZEITRAUMS
    if start_date is None:
        start_date = (datetime.date.today() - datetime.timedelta(days=30)).strftime('%Y-%m-%d')
    if end_date is None:
        end_date = (datetime.date.today() + datetime.timedelta(days=30)).strftime('%Y-%m-%d')
        
    params = {
        'startDate': start_date,
        'endDate': end_date
    }
    
    try:
        response = requests.get(base_url, headers=headers, params=params)
        response.raise_for_status()  
        
        data = response.json()
        
        matching_events = []
        term = search_term.lower()

        for event in data.get('events', []):
            title = event.get('title', '')
            notes = event.get('notes', '')
            is_match = False

            if mode == 'gemeinde':
                # Suche nach Nachname UND (1x, 2x, 3x, 4x, 5x im Titel)
                t_low = title.lower()
                has_name = term in t_low or term in notes.lower()
                has_x = any(x in t_low for x in ["1x", "2x", "3x", "4x", "5x"])
                
                if has_name and has_x:
                    is_match = True
            else:
                if term in title.lower() or term in notes.lower():
                    is_match = True
            
            if is_match:
                
                start_iso = event.get('start_dt')
                end_iso = event.get('end_dt')
                
                if start_iso and end_iso:
                    start_iso_clean = start_iso.replace('Z', '+00:00')
                    end_iso_clean = end_iso.replace('Z', '+00:00')
                    
                    start_dt = datetime.datetime.fromisoformat(start_iso_clean)
                    end_dt = datetime.datetime.fromisoformat(end_iso_clean)
                    
                    event_tuple = (
                        title, 
                        start_dt.strftime('%d.%m.%Y'),
                        start_dt.strftime('%H:%M'),
                        end_dt.strftime('%H:%M')
                    )
                    matching_events.append(event_tuple)
                    
        return matching_events
        
    except requests.exceptions.HTTPError as e:
        error_details = response.text if hasattr(response, 'text') else str(e)
        messagebox.showerror("API Fehler", f"HTTP-Fehler beim Abruf der Teamup-Daten: {e}\nDetails: {error_details}")
        return []
    except Exception as e:
        messagebox.showerror("Fehler", f"Teamup API-Fehler: {e}")
        return []

def fill_template(patient_id, patient_data_tuple, template_data, add_gemeinde_block, ausstellungs_datum, ueberweisung=True): 
    """Füllt die Word-Vorlage mit den Patientendaten und Leistungen und speichert sie."""
    
    # Imports müssen am Anfang der Datei sein, aber wir stellen sicher, dass Pt verfügbar ist
    from docx.shared import Pt 
    from docx.enum.text import WD_UNDERLINE 
    import os
    import datetime
    from docx import Document
    
    # --- Einrückungs-Konstanten ---
    LEISTUNG_INDENT_SIZE = Pt(70)  
    SPACE_AFTER_LEISTUNG_BLOCK = Pt(12) 

    # ... (Datenextraktion bleibt unverändert) ...
    # Unpacking für 13 Elemente (ohne ueberweisung in DB)
    _, vorname, nachname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, _ = patient_data_tuple
    
    # Annahme: get_patient_leistungen_for_template, TEMPLATE_FILE, OUTPUT_FOLDER sind hier verfügbar
    leistungen_liste = get_patient_leistungen_for_template(patient_id)
    
    heute = ausstellungs_datum # Verwendet das übergebene Datum
    
    # Sicherheitscheck: Falls ueberweisung None ist, Standard auf True setzen
    if ueberweisung is None:
        ueberweisung = True
    logging.info(f"Generiere Rechnung. Überweisung-Modus: {ueberweisung}")

    try:
        document = Document(TEMPLATE_FILE)
    except FileNotFoundError:
        raise FileNotFoundError(f"Die Vorlagendatei '{TEMPLATE_FILE}' wurde nicht gefunden.")
    
    total_betrag = 0.0
    
    if not document.paragraphs:
        raise ValueError("Word-Vorlage enthält keine Paragraphen.")

    invoice_number_paragraph = document.paragraphs[0]
            
    # Statische Platzhalter
    replacements = {
        '{{Rechnungsnummer}}': template_data['BHAG_NUMMER'], 
        '{{Anrede}}': anrede, 
        '{{Nachname}}': nachname,
        '{{Vorname}}': vorname,
        '{{Straße}}': strasse,
        '{{Hausnummer}}': hausnummer,
        '{{Adresszusatz}}': adresszusatz or '',
        '{{Postleitzahl}}': plz,
        '{{Stadt}}': ort,
        '{{Versicherungsnummer}}': versicherungsnummer,
        '{{Datum_Austellung}}': heute,
        '{{Diagnose}}': diagnose
    }
    
    # --- 0. Globale Formatierung (Vorbereitung und HONORARNOTE Sonderfall) ---
    
    HONORARNOTE_PARAGRAPH = None
    ORT_DATUM_PARAGRAPH = None
    DIAGNOSE_PLACEHOLDER_PARAGRAPH = None 
    
    for i, p in enumerate(document.paragraphs):
        
        if 'H O N O R A R N O T E' in p.text:
            HONORARNOTE_PARAGRAPH = p
            
        if 'Mödling, {{Datum_Austellung}}' in p.text:
            ORT_DATUM_PARAGRAPH = p
            
        if '{{Diagnose}}' in p.text:
            DIAGNOSE_PLACEHOLDER_PARAGRAPH = p
            
        # Setze die Standardgröße 12pt global
        for run in p.runs:
            run.font.size = Pt(12) 
            
    # Spezielle Formatierung für HONORARNOTE anwenden
    if HONORARNOTE_PARAGRAPH:
        for run in HONORARNOTE_PARAGRAPH.runs:
            run.font.size = Pt(18)
        HONORARNOTE_PARAGRAPH.paragraph_format.space_after = Pt(10) 
        HONORARNOTE_PARAGRAPH.paragraph_format.space_before = Pt(0) 
        HONORARNOTE_PARAGRAPH.paragraph_format.line_spacing = 1.0 


    # --- 1. Rechnungsnummer ersetzen und Gemeinde-Block einfügen ---
    
    if '{{Rechnungsnummer}}' in invoice_number_paragraph.text:
        invoice_number_paragraph.text = invoice_number_paragraph.text.replace('{{Rechnungsnummer}}', template_data['BHAG_NUMMER'])
        del replacements['{{Rechnungsnummer}}']
        
    if add_gemeinde_block:
        # 1. Ziel-Index festlegen (direkt unter der Rechnungsnummer)
        target_index = 1 
        try:
            target_p = document.paragraphs[target_index]
        except IndexError:
            target_p = document.paragraphs[-1]

        # 2. Einen neuen Absatz genau dort einfügen
        new_p = target_p.insert_paragraph_before()
        new_p.paragraph_format.line_spacing = 1.0
        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.space_after = Pt(0)

        # 3. Den Inhalt zusammenbauen:
        # Eine Leerzeile oben (\n)
        # Dann die Adresse
        # Dann vier Leerzeilen unten (\n\n\n\n)
        full_text = (
            "\n\n\n" +
            "Gemeinde Wiener Neudorf\n" +
            "Europaplatz 2\n" +
            "2351 Wiener Neudorf" +
            "\n\n\n\n\n"
        )

        # 4. Den Text dem Absatz hinzufügen und fett/12pt machen
        run = new_p.add_run(full_text)
        run.bold = True
        run.font.size = Pt(12)

    # --- 2. Dynamische Leistungsblock-Logik und restliche Ersetzung (mit Seitenumbruch-Schutz) ---
    
    start_tag = '{{LEISTUNGSBLOCK_START}}'
    end_tag = '{{LEISTUNGSBLOCK_ENDE}}'
    
    block_start_paragraph = None
    block_end_paragraph = None
    block_paragraphs = []
    
    in_block = False 
    
    # Text-Ersetzung und Block-Suche in einem Durchlauf
    for p in document.paragraphs:
        
        if p is HONORARNOTE_PARAGRAPH or p is invoice_number_paragraph or p is DIAGNOSE_PLACEHOLDER_PARAGRAPH:
            continue
            
        # Statische Platzhalter ersetzen
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
        
        # Block-Logik: Musterblock finden
        if start_tag in p.text:
            block_start_paragraph = p
            in_block = True
            continue
        
        if end_tag in p.text:
            block_end_paragraph = p
            break
            
        if in_block:
            block_paragraphs.append(p)

    # Generieren des gesamten Leistungs-Textes
    gesamt_leistungs_text = ""
    
    if block_start_paragraph and block_end_paragraph:
        
        template_text = '\n'.join([p.text for p in block_paragraphs])
        
        for datum_db, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag in leistungen_liste: 
            
            datum_formatiert = datetime.datetime.strptime(datum_db, '%Y-%m-%d').strftime('%d.%m.%Y')
            datum_uhrzeit_text = f"{datum_formatiert}, von {uhrzeit_von} bis {uhrzeit_bis}" 
            
            summe_leistung = einzelbetrag 
            total_betrag += summe_leistung
            
            full_description = beschreibung
            if ' - ' in full_description:
                 full_description = full_description.split(' - ', 1)[-1] 
            
            leistung_block = template_text.replace('{{LEISTUNG_DATUM}}', datum_uhrzeit_text)
            leistung_block = leistung_block.replace('{{LEISTUNG_BESCHREIBUNG}}', full_description) 
            leistung_block = leistung_block.replace('{{LEISTUNG_SUMME}}', f"€ {summe_leistung:.2f}")
            
            # Trennlinie hinzufügen, um später den Abstand zu setzen
            leistung_block += '[[ENDE_BLOCK]]' + '\n\n' 
            
            gesamt_leistungs_text += leistung_block 
            
        if gesamt_leistungs_text:
            
            current_line_is_leistung = False
            lines_of_last_block = [] # Speichert die Zeilen des letzten vollständigen Blocks
            
            for zeile in gesamt_leistungs_text.split('\n'):
                content = zeile.strip() 
                
                if not content:
                    continue

                is_end_block_marker = '[[ENDE_BLOCK]]' in content
                
                # Wir sammeln die Zeilen des aktuellen Blocks, um später keep_with_next zu setzen
                if not is_end_block_marker:
                    lines_of_last_block.append(content)


                if content.startswith("Leistung:"):
                    current_line_is_leistung = True
                elif content.startswith("Datum:") or content.startswith("Betrag:"):
                    current_line_is_leistung = False
                
                
                # Fügen Sie den Paragraphen ein
                if content:
                    
                    # Wenn wir den Ende-Marker des Blocks sehen, ignorieren wir ihn für den Inhalt
                    if is_end_block_marker:
                        content = content.replace('[[ENDE_BLOCK]]', '').strip()
                        if not content:
                            continue

                    p_new = block_start_paragraph.insert_paragraph_before(content)
                    
                    # Allgemeine Formatierung
                    p_new.paragraph_format.space_before = Pt(0)
                    p_new.paragraph_format.space_after = Pt(0)
                    p_new.paragraph_format.line_spacing = 1.0
                    
                    for run in p_new.runs:
                        run.font.size = Pt(12)
                    
                    # Korrektur für Leistungs-Einzug (hängender Einzug)
                    if current_line_is_leistung:
                        p_new.paragraph_format.left_indent = LEISTUNG_INDENT_SIZE      
                        p_new.paragraph_format.first_line_indent = -LEISTUNG_INDENT_SIZE 

                    
                    # 💥 NEUE LOGIK FÜR SEITENUMBRUCH-SCHUTZ:
                    # Der Schutz muss für alle Zeilen eines Blocks aktiviert werden, AUSSER für die letzte Zeile.
                    if not is_end_block_marker:
                        # Setze Keep_with_next für alle Zeilen, die nicht die letzte Zeile des Blocks sind.
                        p_new.paragraph_format.keep_with_next = True
                    else:
                        # Für die letzte Zeile (mit dem ENDE_BLOCK Marker) setzen wir es auf False
                        p_new.paragraph_format.keep_with_next = False
                        
                        # Setze Abstand nach dem Leistungsblock (nach "Betrag:")
                        p_new.paragraph_format.space_after = SPACE_AFTER_LEISTUNG_BLOCK
                        
                        # Setze den Block zurück für den nächsten Durchgang
                        lines_of_last_block = []


            # Entferne die alten Platzhalter-Paragraphen
            block_start_paragraph.text = '' 
            for p in block_paragraphs:
                p._element.getparent().remove(p._element)
            block_end_paragraph.text = ''
        else:
             block_start_paragraph.text = '' 
             for p in block_paragraphs:
                 p._element.getparent().remove(p._element)
             block_end_paragraph.text = ''


    # Ersetzen des Gesamtbetrags
    total_betrag_str = f"{total_betrag:.2f}"
    
    for p in document.paragraphs:
        
        if '{{Gesamt_Betrag}}' in p.text:
            
            original_text = p.text
            p.text = '' 
            
            parts = original_text.split('{{Gesamt_Betrag}}')
            
            run_label = p.add_run(parts[0]) 
            run_label.bold = True
            run_label.font.size = Pt(12)
            
            run_value = p.add_run(total_betrag_str) 
            run_value.bold = True
            run_value.font.underline = WD_UNDERLINE.DOUBLE 
            run_value.font.size = Pt(12)
            
            if len(parts) > 1 and parts[1]:
                run_after = p.add_run(parts[1])
                run_after.bold = True 
                run_after.font.size = Pt(12)
        
        elif 'Gesamt:' in p.text:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)

    
    # --- 3. Letzte Zeilenabstands-Korrektur (Finaler Sweep) ---
    for p in document.paragraphs:
        if p is not HONORARNOTE_PARAGRAPH:
            p.paragraph_format.space_before = Pt(0)
            if p.paragraph_format.space_after is None: 
                p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    # --- 3b. Überweisungstext Logik (NEU) ---
    # Wenn Überweisung = NEIN (0), ersetze den Bank-Block durch "Betrag dankend erhalten!"
    if not ueberweisung:
        for p in document.paragraphs:
            # Suche nach dem Standard-Text aus der Vorlage
            if "Bitte überweisen Sie den Betrag" in p.text:
                p.text = "" # Text löschen
                run = p.add_run("Betrag dankend erhalten!")
                run.bold = True
                run.font.size = Pt(12)
            # Falls IBAN/BIC in eigenen Paragraphen stehen oder Reste davon da sind
            if "BIC GIBA" in p.text or "IBAN AT73" in p.text:
                p.text = ""


    # --- 4. Leerzeilen an den gewünschten Stellen einfügen (Post-Processing) ---
    
    def insert_empty_line(target_p):
        """Fügt eine leere Zeile (Paragraph) VOR dem Ziel-Paragraph ein."""
        # Korrigierte Logik, da insert_paragraph_after nicht existiert
        try:
            p_new = target_p.insert_paragraph_before('')
            
            p_new.paragraph_format.space_before = Pt(0)
            p_new.paragraph_format.space_after = Pt(0)
            p_new.paragraph_format.line_spacing = 1.0
            p_new.add_run('').font.size = Pt(12) 
            
        except (ValueError, IndexError, AttributeError):
            pass

    # Wichtig: Die Leerzeilen werden jetzt VOR dem Paragraphen eingefügt.
    # Wenn Sie eine Leerzeile NACH dem Paragraphen wünschen, müssen Sie
    # den nächsten Paragraphen finden und DIESEM die Leerzeile VORANSTELLEN.
    
    # Hier verwenden wir die einfache Methode, die Leerzeile VOR dem Block einzufügen
    if DIAGNOSE_PLACEHOLDER_PARAGRAPH:
        insert_empty_line(DIAGNOSE_PLACEHOLDER_PARAGRAPH)
        
    if HONORARNOTE_PARAGRAPH:
        # Fügt eine Leerzeile VOR der Honorarnote ein (was nicht gewünscht ist).
        # Lassen Sie dies am besten weg, da die Abstände (space_after) am Anfang
        # des Dokuments besser gesteuert werden sollten.
        pass # insert_empty_line(HONORARNOTE_PARAGRAPH)
        
    if ORT_DATUM_PARAGRAPH:
        insert_empty_line(ORT_DATUM_PARAGRAPH)

    
    # 💥 --- 5. Final Diagnosis Format Override (DIE ULTIMATIVE KORREKTUR!) --- 💥
    if DIAGNOSE_PLACEHOLDER_PARAGRAPH:
        
        target_p = DIAGNOSE_PLACEHOLDER_PARAGRAPH
        p_element = target_p._element
        
        diagnosis_label = "Diagnose:   "
        diagnosis_value = replacements.get('{{Diagnose}}', diagnose)
        
        # 1. Neuen Paragraphen VOR dem alten Platzhalter einfügen
        p_new = target_p.insert_paragraph_before('')
        
        # 2. Formatierung für den Paragraphen festlegen (Abstände und neutraler Stil)
        p_new.paragraph_format.left_indent = Pt(0)      
        p_new.paragraph_format.first_line_indent = Pt(0) 
        p_new.paragraph_format.space_before = Pt(0)
        p_new.paragraph_format.space_after = Pt(0)
        p_new.paragraph_format.line_spacing = 1.0
        
        # 3. Text in zwei Runs einfügen und formatieren
        
        # Run für das Label ("Diagnose: ")
        run_label = p_new.add_run(diagnosis_label)
        run_label.bold = True
        run_label.font.size = Pt(12)
        
        # Run für den Wert (die Diagnose)
        run_value = p_new.add_run(diagnosis_value)
        run_value.bold = True 
        run_value.font.size = Pt(12) 
        
        # 4. Den alten Platzhalter-Paragraphen KOMPLETT aus dem Dokument entfernen
        p_element.getparent().remove(p_element)
        
        # Setze den Platzhalter auf den neuen Paragraphen, um die Leerzeile in Schritt 4 zu erhalten
        DIAGNOSE_PLACEHOLDER_PARAGRAPH = p_new
        
    # --- 6. Footer Zusammenhalten (Keep With Next) ---
    # Verhindert, dass der Block ab "Gesamt:" durch einen Seitenumbruch getrennt wird.
    footer_started = False
    for i, p in enumerate(document.paragraphs):
        if 'Gesamt:' in p.text:
            footer_started = True
        
        if footer_started:
            # keep_with_next für alle Zeilen bis zur vorletzten setzen
            if i < len(document.paragraphs) - 1:
                p.paragraph_format.keep_with_next = True


    # --- Speichern des Dokuments ---
    patient_folder_name = f"{nachname} {vorname}"
    
    output_folder = os.path.expanduser(CONFIG.get('PATIENT_BASE_DIR'))
    patient_output_path = os.path.join(output_folder, patient_folder_name)
    os.makedirs(patient_output_path, exist_ok=True)
    
    output_filename = f"Honorarnote Krankenkasse {template_data['BHAG_NUMMER']}.docx"
    output_path = os.path.join(patient_output_path, output_filename)

    document.save(output_path)
    return output_path

class HonorarGeneratorApp:
    def __init__(self, root):
        self.root = root
        self.root = root
        self.root.title("LeprendiX")
        # Fenster zentrieren + etwas nach rechts versetzt (Platz für Status-Checker links)
        w, h = 1050, 780
        ws = self.root.winfo_screenwidth()
        hs = self.root.winfo_screenheight()
        x = int((ws/2) - (w/2)) + 100
        y = int((hs/2) - (h/2))
        self.root.geometry(f"{w}x{h}+{x}+{y}")
                
        self.patient_data = None  
        self.stammdaten_betraege = {} 
        self.selected_leistung_id = None 
        self.selected_leistungs_kurznamen = set() # Für die Mehrfachauswahl-Buttons
        self.session_blacklist = set() # In-memory blacklist for the current session
        self.ttk_style = ttk.Style(master=root) 
        self.mobile_server = None # Server instance
        
        # NEU: Initialisierung der Folgenummer (BHAG-Logik)
        self.invoice_seq_var = tk.StringVar(master=root) 
        self.invoice_sequence_data = self._get_invoice_sequence_data()
        self.invoice_seq_var.set(str(self.invoice_sequence_data.get('rechnung_folgenummer', '0')).zfill(3))

        # --- QoL: Status Bar ---
        self.status_var = tk.StringVar()
        self.status_bar = ttk.Label(root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
        
        self.add_gemeinde_block_var = tk.BooleanVar(master=root, value=0) # Standardmäßig AUS
        self.use_km_money_var = tk.BooleanVar(master=root, value=True) # NEU: Standardmäßig AN
        self.gen_ueberweisung_var = tk.BooleanVar(master=root, value=True) # Standardmäßig AN (Überweisung)

        # In HonorarGeneratorApp.__init__ (nach self.add_gemeinde_block_var = ...)
        now = datetime.datetime.now()
        self.selected_invoice_month = tk.StringVar(master=root, value=f"{now.month:02d}")
        self.selected_invoice_year = tk.StringVar(master=root, value=str(now.year))

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(pady=10, padx=10, expand=True, fill="both")

        self.ttk_style = ttk.Style(master=root) 
        # NEU: Der Stil MUSS ebenfalls HIER gesetzt werden, da er in setup_patient_tab verwendet wird.
        self.ttk_style.configure('Danger.TButton', foreground='red')

        self.tab_generate = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_generate, text='📝 Honorarnote Generieren')
        self.setup_generate_tab(self.tab_generate)
        
        self.tab_patient = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_patient, text='👤 Patienten Verwalten')
        self.setup_patient_tab(self.tab_patient)

        self.tab_leistung = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_leistung, text='➕ Leistungen Hinzufügen/Prüfen')
        self.setup_leistung_tab(self.tab_leistung)
        self.update_leistung_list()
        
        self.tab_stammdaten = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_stammdaten, text='⚙️ Stammdaten Leistungen')
        self.setup_stammdaten_tab(self.tab_stammdaten)
        

        self.add_gemeinde_block_var.set(0)

              
        self.update_patient_info() 
        self.load_leistung_stammdaten_buttons()
        self.open_status_checker()
        self.focus_and_highlight(self.search_entry)
        
        # Hotkeys aus Config laden und binden
        hk_enter = CONFIG.get('HOTKEY_ENTER', '<Return>')
        hk_switch = CONFIG.get('HOTKEY_SWITCH_TAB', '<F12>, <Delete>')

        try:
            if hk_enter:
                self.root.bind(hk_enter, self.handle_global_enter)
        except Exception as e:
            logging.error(f"Fehler beim Binden von Enter-Hotkey: {e}")

        if hk_switch:
            for k in hk_switch.split(','):
                if k.strip():
                    try:
                        self.root.bind(k.strip(), self._switch_to_generate_tab)
                    except Exception as e:
                        logging.error(f"Fehler beim Binden von Switch-Hotkey '{k}': {e}")
        
        # NEU: Kontextmenü für Textfelder einrichten
        self._setup_context_menu()

        self.root.protocol("WM_DELETE_WINDOW", self.on_closing)

    def on_closing(self):
        if messagebox.askyesno("Backup", "Möchten Sie vor dem Beenden ein automatisches Backup erstellen?", parent=self.root):
            if os.path.exists(DATABASE_NAME):
                try:
                    base_dir = os.path.dirname(DATABASE_NAME)
                    backup_dir = os.path.join(base_dir, "backups")
                    os.makedirs(backup_dir, exist_ok=True)
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    backup_path = os.path.join(backup_dir, f"autobackup_{timestamp}.db")
                    shutil.copy2(DATABASE_NAME, backup_path)
                    logging.info(f"[AutoBackup] Backup erstellt: {backup_path}")
                except Exception as e:
                    logging.error(f"[AutoBackup] Fehler: {e}")
        self.root.destroy()

    def _setup_context_menu(self):
        """Bindet das Rechtsklick-Menü an alle Text-Widgets."""
        for widget_class in ["Entry", "TEntry", "Text"]:
            self.root.bind_class(widget_class, "<Button-3>", self._show_context_menu)

    def refresh_all_stammdaten_ui(self):
        """Refreshes all UI components that display master data for services."""
        logging.info("Refreshing Stammdaten UI...")
        try:
            self.update_stammdaten_list()
            self.load_leistung_stammdaten_buttons()
            self.set_status("Stammdaten-Anzeige wurde aktualisiert.", 2000)
        except Exception as e:
            logging.error(f"Error refreshing stammdaten UI: {e}")
            messagebox.showerror("Fehler", f"Fehler beim Aktualisieren der Stammdaten-Anzeige: {e}")

    def _show_context_menu(self, event):
        """Zeigt ein Kontextmenü mit Kopieren/Einfügen."""
        try:
            event.widget.focus_set()
            menu = tk.Menu(self.root, tearoff=0)
            menu.add_command(label="Kopieren", command=lambda: event.widget.event_generate("<<Copy>>"))
            menu.add_command(label="Einfügen", command=lambda: event.widget.event_generate("<<Paste>>"))
            menu.tk_popup(event.x_root, event.y_root)
        except Exception:
            pass

    def set_status(self, message, duration=4000):
        """Setzt eine Nachricht in der Statusleiste, die nach 'duration' ms verschwindet."""
        self.status_var.set(f" {message}")
        if duration:
            self.root.after(duration, lambda: self.status_var.set(""))

    def _blink_highlight(self, widget, count=6):
        if not widget or not widget.winfo_exists() or count <= 0:
            # Restore original state at the end
            if hasattr(widget, '_original_style'):
                try:
                    widget.configure(style=widget._original_style)
                except tk.TclError:
                    pass
                del widget._original_style
            if hasattr(widget, '_original_bg'):
                try:
                    widget.configure(background=widget._original_bg)
                except tk.TclError:
                    pass
                del widget._original_bg
            return

        is_on = count % 2 == 0
        
        if is_on:
            # --- Turn ON highlight ---
            if isinstance(widget, (ttk.Entry, ttk.Button)):
                style_name = widget.winfo_class()
                highlight_style_name = f"Highlight.{style_name}"
                
                if not hasattr(widget, '_original_style'):
                    widget._original_style = widget.cget("style") or style_name

                s = self.ttk_style
                try:
                    s.layout(highlight_style_name)
                except tk.TclError: # Style does not exist, create it
                    if isinstance(widget, ttk.Entry):
                        s.configure(highlight_style_name, fieldbackground='#3498db')
                    elif isinstance(widget, ttk.Button):
                        s.configure(highlight_style_name, background='#3498db', foreground='white')
                        try:
                            original_font = s.lookup(widget._original_style, 'font')
                            if original_font:
                                s.configure(highlight_style_name, font=original_font)
                        except tk.TclError:
                            pass # Use default font
                widget.configure(style=highlight_style_name)
            elif isinstance(widget, tk.Listbox):
                if not hasattr(widget, '_original_bg'):
                    widget._original_bg = widget.cget("background")
                widget.configure(background='#3498db')
        else:
            # --- Turn OFF highlight ---
            if hasattr(widget, '_original_style'): widget.configure(style=widget._original_style)
            if hasattr(widget, '_original_bg'): widget.configure(background=widget._original_bg)

        self.root.after(150, lambda: self._blink_highlight(widget, count - 1))

    def focus_and_highlight(self, widget):
        if widget and widget.winfo_exists():
            widget.focus_set()
            self._blink_highlight(widget)
   
    def _get_invoice_sequence_data(self):
        """Holt die aktuelle Folgenummer, Jahr und Monat aus der Einstellungen-Tabelle."""
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        data = {}
        now = datetime.datetime.now()
        
        # Standardwerte (für den Fall, dass die DB-Tabelle noch nicht existiert)
        default_data = {
            'rechnung_jahr': str(now.year),
            'rechnung_monat': str(now.month).zfill(2),
            'rechnung_folgenummer': '0', # Start bei 0
            'last_invoice_path': '',
            'last_invoice_patient_id': ''
        }
        
        try:
            cursor.execute("SELECT key, value FROM einstellungen WHERE key IN ('rechnung_jahr', 'rechnung_monat', 'rechnung_folgenummer', 'last_invoice_path', 'last_invoice_patient_id')")
            for key, value in cursor.fetchall():
                data[key] = value
            
            # Stelle sicher, dass alle Schlüssel vorhanden sind
            return {**default_data, **data}
            
        except sqlite3.OperationalError:
            # Fallback, wenn die Tabelle noch nicht existiert (z.B. wenn db_setup.py nicht lief)
            print("WARNUNG: Einstellungen-Tabelle nicht gefunden. Verwende Standardwerte.")
            return default_data
        finally:
            conn.close()

    def _update_invoice_sequence_data(self, key, value):
        """Aktualisiert oder fügt einen Wert in der Einstellungen-Tabelle ein."""
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        try:
            # Versucht, den Wert zu aktualisieren (UPDATE OR IGNORE ist sicherer)
            cursor.execute("UPDATE einstellungen SET value = ? WHERE key = ?", (str(value), key))
            if cursor.rowcount == 0:
                 # Wenn kein Update möglich war, wird ein INSERT versucht
                 cursor.execute("INSERT OR IGNORE INTO einstellungen (key, value) VALUES (?, ?)", (key, str(value)))
            conn.commit()
        except Exception as e:
            logging.error(f"FEHLER beim Aktualisieren der Folgenummer in der DB ({key}): {e}")
            messagebox.showerror("Fehler", f"Fehler beim Aktualisieren der Folgenummer in der DB ({key}): {e}")
        finally:
            conn.close()

    def _save_custom_invoice_number(self):
        """Speichert die manuell eingegebene Folgenummer und aktualisiert die DB."""
        try:
            new_nummer_str = self.invoice_seq_var.get().strip()
            
            # Validierung: Muss eine Zahl sein (max. 3 Ziffern, wir akzeptieren bis 999)
            if not new_nummer_str.isdigit() or len(new_nummer_str) > 3 or int(new_nummer_str) < 0:
                messagebox.showerror("Fehler", "Folgenummer muss eine Zahl zwischen 0 und 999 sein.")
                # Setzt auf den letzten gültigen Wert zurück (intern gespeichert)
                self.invoice_seq_var.set(str(self.invoice_sequence_data['rechnung_folgenummer']).zfill(3)) 
                return

            new_nummer = int(new_nummer_str)

            # Aktualisiere interne Daten und DB
            self.invoice_sequence_data['rechnung_folgenummer'] = str(new_nummer)
            self._update_invoice_sequence_data('rechnung_folgenummer', str(new_nummer))
            
            # Aktualisiere das Anzeige-Feld (mit führenden Nullen)
            self.invoice_seq_var.set(str(new_nummer).zfill(3))
            
            messagebox.showinfo("Erfolg", f"Folgenummer manuell auf '{str(new_nummer).zfill(3)}' gespeichert.")
            
        except Exception as e:
            messagebox.showerror("Fehler", f"Fehler beim Speichern der Folgenummer: {e}")

    def _prepare_bhag_number(self):
        """Generiert die nächste BHAG-Nummer, aktualisiert die DB und die GUI."""
        logging.debug(f"Sequenz-Daten aktuell: {self.invoice_sequence_data}")


        now = datetime.datetime.now()
        current_year = str(now.year)
        current_month = str(now.month).zfill(2) 

        # Hole die gespeicherten Daten (als String)
        stored_year = self.invoice_sequence_data.get('rechnung_jahr', current_year)
        stored_month = self.invoice_sequence_data.get('rechnung_monat', current_month)
        current_folgenummer = int(self.invoice_sequence_data.get('rechnung_folgenummer', '0'))

        # --- Reset-Logik (Jährlich) ---
        if current_year != stored_year:
            # Jährlicher Reset: Setze auf 1 und aktualisiere Jahr/Monat
            new_folgenummer = 1
            self.invoice_sequence_data['rechnung_jahr'] = current_year
            self.invoice_sequence_data['rechnung_monat'] = current_month
            self._update_invoice_sequence_data('rechnung_jahr', current_year)
            self._update_invoice_sequence_data('rechnung_monat', current_month)
        else:
            # Normales Inkremement
            new_folgenummer = current_folgenummer + 1

        # 2. Folgenummer formatieren und BHAG-Nummer generieren
        folgenummer_str = str(new_folgenummer).zfill(3) # z.B. '001', '010'
        self.invoice_seq_var.set(folgenummer_str)
        self.invoice_sequence_data['rechnung_folgenummer'] = folgenummer_str
        bhag_nummer = f"BHAG{current_year}{current_month}{folgenummer_str}"

        # 3. Datenbank und GUI aktualisieren
        self.invoice_sequence_data['rechnung_folgenummer'] = str(new_folgenummer)
        self._update_invoice_sequence_data('rechnung_folgenummer', str(new_folgenummer))
        self.invoice_seq_var.set(folgenummer_str) # Aktualisiere das Feld im GUI
        logging.info(f"Generierte Nummer: {bhag_nummer}")
        
        # Stelle sicher, dass die Monatsangabe immer aktuell ist (falls sie sich innerhalb des Jahres ändert)
        if stored_month != current_month:
            self.invoice_sequence_data['rechnung_monat'] = current_month
            self._update_invoice_sequence_data('rechnung_monat', current_month)

        return {'BHAG_NUMMER': bhag_nummer}
        
    # --- 1. Generieren Tab ---
    def setup_generate_tab(self, tab):
        now = datetime.datetime.now()
        ttk.Label(tab, text="Patienten-Suche (Nachname/Vorname):").grid(row=0, column=0, padx=5, pady=5, sticky='w')
        self.search_entry = ttk.Entry(tab, width=40)
        self.search_entry.grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Button(tab, text="Suchen", command=self.search_patients).grid(row=0, column=2, padx=5, pady=5)
        # NEU: Button zum Öffnen des Status-Checkers
        ttk.Button(tab, text="Status-Checker", command=self.open_status_checker).grid(row=0, column=3, padx=5, pady=5)
        
        self.results_listbox = tk.Listbox(tab, height=10, width=60)
        self.results_listbox.grid(row=1, column=0, columnspan=3, padx=5, pady=5, sticky='nsew')
        self.results_listbox.bind('<<ListboxSelect>>', self.select_patient_from_list)

        self.search_entry.bind('<Return>', lambda e: self.search_patients()) # Enter 1: Suchen
        self.results_listbox.bind('<Return>', self.handle_enter_on_listbox) # Enter 2: Patient wählen


        ttk.Label(tab, text="Aktueller Patient:").grid(row=2, column=0, padx=5, pady=5, sticky='w')
        self.current_patient_label = ttk.Label(tab, text="Kein Patient ausgewählt", foreground='blue')
        self.current_patient_label.grid(row=2, column=1, columnspan=2, padx=5, pady=5, sticky='w')

        # Frame für die Generierungs-Buttons
        btn_frame = ttk.Frame(tab)
        btn_frame.grid(row=3, column=0, columnspan=3, pady=20)
        
        # Bestehender Button
        ttk.Button(btn_frame, text="Speichern & Öffnen", command=self.generate_invoice).pack(side=tk.LEFT, padx=10)
        
        # NEUER BUTTON: Generieren und Sofort Drucken
        ttk.Button(btn_frame, text="✅ Speichern & Drucken", command=self.generate_and_print_invoice).pack(side=tk.LEFT, padx=10)
        
        # NEU: Button zum Widerrufen der letzten Honorarnote
        ttk.Button(btn_frame, text="Honorarnote Wiederrufen", command=self.revoke_last_invoice, style='Danger.TButton').pack(side=tk.LEFT, padx=10)
        
        # --- Honorarnoten-Folgenummer (NEU) ---
        folgenummer_frame = ttk.LabelFrame(tab, text="Honorar-Folgenummer (BHAG-Nr.)")
        folgenummer_frame.grid(row=4, column=0, columnspan=3, padx=10, pady=5, sticky="ew")

        ttk.Label(folgenummer_frame, text="Aktuelle Folgenummer (000-999):").grid(row=0, column=0, padx=5, pady=2, sticky="w")
        self.invoice_seq_entry = ttk.Entry(folgenummer_frame, textvariable=self.invoice_seq_var, width=5)
        self.invoice_seq_entry.grid(row=0, column=1, padx=5, pady=2, sticky="w")

        save_folgenummer_btn = ttk.Button(folgenummer_frame, text="Folgenummer speichern (manuell korrigieren)", command=self._save_custom_invoice_number)
        save_folgenummer_btn.grid(row=0, column=2, padx=5, pady=2, sticky="w")

        ttk.Label(folgenummer_frame, text="Format: BHAG[Jahr][Monat][Folgenummer]").grid(row=1, column=0, columnspan=3, padx=5, pady=2, sticky="w")

        zusatz_frame = ttk.LabelFrame(tab, text="Zusätzliche Optionen für Word-Generierung")
        zusatz_frame.grid(row=5, column=0, columnspan=3, padx=10, pady=5, sticky="ew")
        ttk.Checkbutton(zusatz_frame, text="Zusätzlichen 'Gemeinde Wiener Neudorf' Block in Dokument einfügen", variable=self.add_gemeinde_block_var).grid(row=0, column=0, padx=5, pady=5, sticky='w')
        
        ttk.Checkbutton(zusatz_frame, text="Überweisung? (Sonst: 'Betrag dankend erhalten!')", variable=self.gen_ueberweisung_var).grid(row=1, column=0, padx=5, pady=5, sticky='w')
        
        # In setup_generate_tab hinzufügen (z.B. nach dem folgenummer_frame):
        date_selection_frame = ttk.LabelFrame(tab, text="Ausstellungsdatum (Letzter Tag des Monats)")
        date_selection_frame.grid(row=6, column=0, columnspan=3, padx=10, pady=5, sticky="ew")

        ttk.Label(date_selection_frame, text="Monat:").grid(row=0, column=0, padx=5, pady=5)
        month_combo = ttk.Combobox(date_selection_frame, textvariable=self.selected_invoice_month, 
                                values=[str(i).zfill(2) for i in range(1, 13)], width=5, state="readonly")
        month_combo.grid(row=0, column=1, padx=5, pady=5)

        ttk.Label(date_selection_frame, text="Jahr:").grid(row=0, column=2, padx=5, pady=5)
        year_combo = ttk.Combobox(date_selection_frame, textvariable=self.selected_invoice_year, 
                              values=[str(y) for y in range(now.year-1, now.year+2)], # Hier wird 'now' verwendet!
                              width=7, state="readonly")
        year_combo.grid(row=0, column=3, padx=5, pady=5)

        ttk.Label(date_selection_frame, text="Info: Es wird automatisch der letzte Tag des gewählten Monats verwendet.").grid(row=1, column=0, columnspan=4, padx=5, pady=2)

        # Wichtig: Den row-Index des nachfolgenden zusatz_frame auf row=6 anpassen!

        tab.grid_rowconfigure(1, weight=1)

    def open_status_checker(self):
        from patient_status_checker import PatientStatusApp
        checker_window = tk.Toplevel(self.root)
        app = PatientStatusApp(checker_window, 
                               selection_callback=self.load_patient_from_checker,
                               archive_callback=self.refresh_all_stammdaten_ui) # Pass callback
        
        # Auto-Positionierung: Links neben dem Hauptfenster
        self.root.update_idletasks()
        x = self.root.winfo_x()
        y = self.root.winfo_y()
        checker_window.geometry(f"+{x - 410}+{y}") # 400 Breite + 10px Abstand
        
    def load_patient_from_checker(self, patient_id):
        # Setze ID in Suchfeld
        self.search_entry.delete(0, tk.END)
        self.search_entry.insert(0, patient_id)
        # Führe Suche aus
        self.search_patients()
        # Wähle automatisch das erste Ergebnis (sollte der Patient sein)
        if self.results_listbox.size() > 0:
            self.results_listbox.selection_clear(0, tk.END)
            self.results_listbox.selection_set(0)
            self.select_patient_from_list()

        # Fokus auf das Hauptfenster setzen
        self.root.lift()
        self.root.focus_force()

    def search_patients(self):
        search_term = self.search_entry.get().strip()
        if not search_term:
            messagebox.showwarning("Suche", "Bitte geben Sie einen Suchbegriff ein.")
            return
        

        query = self.search_entry.get().strip()
        if not query:
            return

        results = get_patient_data(search_term)
        
        self.results_listbox.delete(0, tk.END)
        self.results_listbox.patient_data_list = []

        if not results:
            self.results_listbox.insert(tk.END, "Keine Patienten gefunden.")
            return

        for patient_data_tuple in results:
            display_text = f"ID {patient_data_tuple[0]} - {patient_data_tuple[2]} {patient_data_tuple[1]} ({patient_data_tuple[7]})"
            self.results_listbox.insert(tk.END, display_text)
            self.results_listbox.patient_data_list.append(patient_data_tuple)
        
        if self.results_listbox.size() > 0:
            self.results_listbox.selection_clear(0, tk.END)
            self.results_listbox.selection_set(0)
            self.results_listbox.activate(0)
            self.focus_and_highlight(self.results_listbox) # WICHTIG: Fokus springt für das nächste ENTER in die Liste


    def finish_import_transition(self):
        """Wechselt sicher zum letzten Tab."""
        self.notebook.select(2) # Zum Drucken-Tab wechseln
        self.root.focus_set()   # Fokus zurück aufs Hauptfenster

    def handle_global_enter(self, event):
        """Zentrale Steuerung der ENTER-Taste."""
        current_tab_index = self.notebook.index(self.notebook.select())
        
        # TAB: Honorarnote generieren (Suche)
        if current_tab_index == 0:
            focused_widget = self.root.focus_get()
            if focused_widget == self.search_entry:
                self.search_patients()
            elif focused_widget == self.results_listbox:
                self.handle_enter_on_listbox(None)
        
        # TAB: Leistungen hinzufügen/prüfen
        elif current_tab_index == 2:
            # Nur öffnen, wenn nicht bereits ein Teamup-Fenster offen ist
            # Wir prüfen spezifisch auf Teamup-Fenster, da andere Toplevels (z.B. Status Checker) existieren können
            teamup_open = False
            for child in self.root.winfo_children():
                if isinstance(child, tk.Toplevel) and child.winfo_exists():
                    try:
                        if "Teamup" in child.title() or "API Abruf" in child.title():
                            teamup_open = True
                            break
                    except Exception:
                        pass
            
            if not teamup_open:
                # Hier rufen wir direkt die Funktion mit Validierung & Ladebalken auf
                if hasattr(self, 'teamup_button'):
                    self.teamup_button.invoke()

    def _get_selected_invoice_date(self):
        year = int(self.selected_invoice_year.get())
        month = int(self.selected_invoice_month.get())
        last_day = calendar.monthrange(year, month)[1]
        return datetime.date(year, month, last_day).strftime("%d.%m.%Y")

    def handle_enter_on_listbox(self, event=None):
        """Verarbeitet die Auswahl eines Patienten per Enter-Taste."""
        selection = self.results_listbox.curselection()
        if selection:
            # Hier wird die Methode nun sicher aufgerufen
            self.select_patient_from_list()

    
    def select_patient_from_list(self, event=None):
        selection = self.results_listbox.curselection()
        if selection:
            index = selection[0]
            self.patient_data = self.results_listbox.patient_data_list[index] 
            self.update_patient_info()
            
            # NEU: Lade und wähle die zuletzt ausgewählten Leistungen aus
            self.load_and_select_last_leistungen()
            
            # Automatisch zum Leistungs-Tab wechseln und Liste aktualisieren
            self.notebook.select(self.tab_leistung)
            self.update_leistung_list() 

            # Focus auf den Teamup-Button setzen
            if hasattr(self, 'teamup_button'):
                self.focus_and_highlight(self.teamup_button)

    def update_patient_info(self):
        if self.patient_data:
            # patient_data[11] ist das Kilometergeld
            km_geld = self.patient_data[11] if len(self.patient_data) > 11 and self.patient_data[11] is not None else 0.0
            text = f"{self.patient_data[2]} {self.patient_data[1]} (ID: {self.patient_data[0]}, VersNr: {self.patient_data[9]}, KM: €{km_geld:.2f})"
            self.current_patient_label.config(text=text)
            if hasattr(self, 'leistung_patient_label'):
                self.leistung_patient_label.config(text=text, foreground='blue')
        else:
            self.current_patient_label.config(text="Kein Patient ausgewählt")
            if hasattr(self, 'leistung_patient_label'):
                self.leistung_patient_label.config(text="Bitte Patient in Tab 1 auswählen", foreground='red')

    def _switch_to_generate_tab(self, event=None): # NEU
        """Behandelt die 'Druck'-Taste je nach aktivem Tab."""
        current_tab = self.notebook.index(self.notebook.select())
        
        # Fall A: Wir sind im Tab "Leistungen hinzufügen" (Index 2)
        # Aktion: Wechsel zum ersten Tab (wie "Fertig -->")
        if current_tab == 2:
            self.notebook.select(0)
            self.focus_and_highlight(self.search_entry)
            
        # Fall B: Wir sind im Tab "Honorarnote generieren" (Index 0)
        # Aktion: Dokument erstellen (wie "Drucken und Speichern")
        elif current_tab == 0:
            # Hier rufen wir die Funktion auf, die dein Button "Drucken und Speichern" nutzt
            # Laut deinem Code-Stil heißt diese vermutlich generate_honorarnote
            self.generate_and_print_invoice()

    def generate_invoice(self):
        """Generiert die Honorarnote und öffnet die Datei."""
        if not self.patient_data:
            messagebox.showwarning("Warnung", "Bitte wählen Sie zuerst einen Patienten aus.")
            return

        if not get_patient_leistungen(self.patient_data[0]):
             messagebox.showwarning("Achtung", "Keine Leistungen für den ausgewählten Patienten gefunden. Druckvorgang abgebrochen.")
             return
        
        template_data = self._prepare_bhag_number() 
        add_gemeinde_block = self.add_gemeinde_block_var.get()
        patient_id = self.patient_data[0]
        ausstellungs_datum = self._get_selected_invoice_date()
        ueberweisung = bool(self.gen_ueberweisung_var.get())
        
        try:
            output_path = fill_template(self.patient_data[0], self.patient_data, template_data, add_gemeinde_block, ausstellungs_datum, ueberweisung=ueberweisung)
            
            # NEU: Pfad und Patient-ID der letzten Rechnung speichern
            self._update_invoice_sequence_data('last_invoice_path', output_path)
            self._update_invoice_sequence_data('last_invoice_patient_id', patient_id)

            # NEU: Setze den Status des Patienten auf "abgerechnet" (1 = Grün)
            _update_invoiced_status(patient_id, 1)

            self.set_status(f"Honorarnote erstellt: {output_path}")
            
            # Öffnen der Datei nach Generierung
            if sys.platform.startswith('win'):
                os.startfile(output_path)
            elif sys.platform.startswith('darwin'):
                subprocess.call(('open', output_path))
            elif sys.platform.startswith('linux'):
                subprocess.call(('xdg-open', output_path))
                
            # TODO: Hier Funktion zum Markieren der Leistungen als "abgerechnet" aufrufen
            
        except FileNotFoundError as e:
             messagebox.showerror("Fehler", str(e))
        except Exception as e:
            import traceback
            messagebox.showerror("Fehler", f"Fehler bei der Generierung: {e}\n\nDetails siehe Konsole.")
            traceback.print_exc()

    def generate_and_print_invoice(self):
        """Generiert die Honorarnote und versucht, sie direkt zu drucken."""
        if not self.patient_data:
            messagebox.showwarning("Warnung", "Bitte wählen Sie zuerst einen Patienten aus.")
            return
        
        if not get_patient_leistungen(self.patient_data[0]):
             messagebox.showwarning("Achtung", "Keine Leistungen für den ausgewählten Patienten gefunden. Druckvorgang abgebrochen.")
             return

        template_data = self._prepare_bhag_number() 
        add_gemeinde_block = self.add_gemeinde_block_var.get()
        ausstellungs_datum = self._get_selected_invoice_date()
        ueberweisung = bool(self.gen_ueberweisung_var.get())

        try:
            # NEU: BHAG-Nummer generieren und DB aktualisieren
            output_path = fill_template(self.patient_data[0], self.patient_data, template_data, add_gemeinde_block, ausstellungs_datum, ueberweisung=ueberweisung)
            add_gemeinde_block = self.add_gemeinde_block_var.get()
            
            # NEU: Pfad und Patient-ID der letzten Rechnung speichern
            self._update_invoice_sequence_data('last_invoice_path', output_path)
            self._update_invoice_sequence_data('last_invoice_patient_id', self.patient_data[0])

            # 2. Versuche, sie sofort zu drucken
            success, message = print_document_silently(output_path)
            
            if success:
                # NEU: Setze den Status des Patienten auf "abgerechnet" (1 = Grün)
                _update_invoiced_status(self.patient_data[0], 1)
                
                self.set_status(f"Honorarnote gedruckt. {message}")
                # TODO: Hier Funktion zum Markieren der Leistungen als "abgerechnet" aufrufen
                
            else:
                # Fallback: Datei anzeigen, falls Drucken fehlschlägt
                if sys.platform.startswith('win'):
                    os.startfile(output_path)
                elif sys.platform.startswith('darwin'):
                    subprocess.call(('open', output_path))
                elif sys.platform.startswith('linux'):
                    subprocess.call(('xdg-open', output_path))
                    
                messagebox.showwarning("Druckfehler", f"Druckauftrag konnte nicht direkt gesendet werden:\n{message}\nDie Datei wurde im Viewer geöffnet.")
                
        except FileNotFoundError as e:
             messagebox.showerror("Fehler", str(e))
        except Exception as e:
            import traceback
            messagebox.showerror("Fehler", f"Fehler bei Generierung/Druck: {e}")
            traceback.print_exc()

    def revoke_last_invoice(self):
        """
        Widerruft die letzte erstellte Honorarnote.
        - Löscht die Datei.
        - Setzt die Rechnungsnummer zurück.
        - Setzt den Patientenstatus auf 'offen' (Rot).
        """
        self.invoice_sequence_data = self._get_invoice_sequence_data()
        last_invoice_path = self.invoice_sequence_data.get('last_invoice_path')
        last_patient_id = self.invoice_sequence_data.get('last_invoice_patient_id')
        
        if not last_invoice_path or not os.path.exists(last_invoice_path):
            messagebox.showwarning("Fehler", "Keine zu widerrufende Honorarnote gefunden oder Datei existiert nicht mehr.")
            return

        if not messagebox.askyesno("Bestätigung", f"Soll die letzte Honorarnote wirklich widerrufen werden?\n\nDatei: {os.path.basename(last_invoice_path)}\n\nDiese Aktion löscht die Datei, setzt die Folgenummer zurück und markiert den Patienten als 'offen'."):
            return

        try:
            # 1. Datei löschen
            os.remove(last_invoice_path)
            logging.info(f"Honorarnote-Datei gelöscht: {last_invoice_path}")

            # 2. Folgenummer dekrementieren
            current_folgenummer = int(self.invoice_sequence_data.get('rechnung_folgenummer', '0'))
            new_folgenummer = max(0, current_folgenummer - 1)
            
            self._update_invoice_sequence_data('rechnung_folgenummer', str(new_folgenummer))
            self.invoice_sequence_data['rechnung_folgenummer'] = str(new_folgenummer)
            self.invoice_seq_var.set(str(new_folgenummer).zfill(3))
            
            # 3. Pfad und Patienten-ID in DB löschen
            self._update_invoice_sequence_data('last_invoice_path', '')
            self._update_invoice_sequence_data('last_invoice_patient_id', '')

            # 4. Patientenstatus zurücksetzen (auf 0 = Rot/Offen)
            if last_patient_id:
                _update_invoiced_status(int(last_patient_id), 0)
                logging.info(f"Patientenstatus für ID {last_patient_id} auf 'offen' zurückgesetzt.")

            messagebox.showinfo("Erfolg", "Die letzte Honorarnote wurde widerrufen.")

        except Exception as e:
            logging.error(f"Fehler beim Widerrufen der Honorarnote: {e}")
            messagebox.showerror("Fehler", f"Ein Fehler ist aufgetreten:\n{e}")


    def _get_ort_for_plz(self, plz):
        """Gibt den Ortsnamen für eine PLZ zurück (Basis-Datenbank für NÖ/Wien)."""
        # Erweiterbare Liste für die Region Mödling/Baden/Wien Umgebung
        mapping = {
            "2351": "Wiener Neudorf", "2340": "Mödling", "2344": "Maria Enzersdorf",
            "2345": "Brunn am Gebirge", "2331": "Vösendorf", "2333": "Leopoldsdorf",
            "2334": "Vösendorf", "2352": "Gumpoldskirchen", "2353": "Guntramsdorf",
            "2361": "Laxenburg", "2362": "Biedermannsdorf", "2380": "Perchtoldsdorf",
            "2384": "Breitenfurt bei Wien", "2391": "Kaltenleutgeben", "2392": "Sulz im Wienerwald",
            "2483": "Ebreichsdorf", "2500": "Baden", "2511": "Pfaffstätten",
            "2512": "Tribuswinkel", "2514": "Traiskirchen", "2521": "Trumau",
            "2522": "Oberwaltersdorf", "2540": "Bad Vöslau", "2542": "Kottingbrunn",
            "2544": "Leobersdorf", "2700": "Wiener Neustadt", "3100": "St. Pölten"
        }
        
        if plz in mapping:
            return mapping[plz]
            
        # Generische Wien-Logik (1010 - 1230)
        if len(plz) == 4 and plz.startswith("1") and plz.endswith("0"):
            try:
                bezirk = int(plz[1:3])
                if 1 <= bezirk <= 23:
                    return "Wien"
            except ValueError:
                pass
        return None

    def autofill_ort(self, event):
        """Füllt das Ort-Feld automatisch basierend auf der PLZ."""
        plz = self.patient_entries["PLZ"].get().strip()
        # Nur reagieren, wenn 4 Zeichen eingegeben wurden
        if len(plz) == 4:
            ort = self._get_ort_for_plz(plz)
            if ort:
                self.patient_entries["Ort"].delete(0, tk.END)
                self.patient_entries["Ort"].insert(0, ort)
                # NEU: Löst die Routenberechnung aus, nachdem der Ort automatisch ausgefüllt wurde
                self._start_distance_calculation()

    def _start_distance_calculation(self, event=None):
        """Startet die Distanzberechnung in einem separaten Thread."""
        self.distance_label.config(text="wird geladen...")
        thread = threading.Thread(target=self._perform_distance_calculation)
        thread.daemon = True
        thread.start()

    def _perform_distance_calculation(self):
        """Berechnet die Routen-Distanz und aktualisiert das GUI-Label."""
        try:
            street = self.patient_entries["Straße"].get().strip()
            housenumber = self.patient_entries["Hausnummer"].get().strip()
            plz = self.patient_entries["PLZ"].get().strip()
            city = self.patient_entries["Ort"].get().strip()

            if not all([street, housenumber, plz, city]):
                self.root.after(0, lambda: self.distance_label.config(text=""))
                return

            patient_address = f"{street} {housenumber}, {plz} {city}, Austria"
            fixed_address = CONFIG.get("FIXED_ADDRESS")

            if not fixed_address:
                self.root.after(0, lambda: self.distance_label.config(text="Fixe Adresse fehlt"))
                return

            geolocator = Nominatim(user_agent="leprendix-distance-calculator")
            
            location1 = geolocator.geocode(patient_address, timeout=10)
            if not location1:
                self.root.after(0, lambda: self.distance_label.config(text="Patienten-Adresse ungültig"))
                return

            location2 = geolocator.geocode(fixed_address, timeout=10)
            if not location2:
                self.root.after(0, lambda: self.distance_label.config(text="Fixe Adresse ungültig"))
                return

            coords1 = (location1.longitude, location1.latitude)
            coords2 = (location2.longitude, location2.latitude)
            
            url = f"http://router.project-osrm.org/route/v1/driving/{coords1[0]},{coords1[1]};{coords2[0]},{coords2[1]}"
            
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            
            data = response.json()
            
            if data['code'] == 'Ok':
                distance_meters = data['routes'][0]['distance']
                distance_km = distance_meters / 1000
                self.root.after(0, lambda: self.distance_label.config(text=f"ca. {distance_km:.1f} km"))
            else:
                self.root.after(0, lambda: self.distance_label.config(text="Route nicht gefunden"))

        except (requests.exceptions.RequestException, GeocoderTimedOut, GeocoderUnavailable) as e:
            self.root.after(0, lambda: self.distance_label.config(text="API Fehler"))
            logging.error(f"Fehler bei der Distanzberechnung: {e}")
        except Exception as e:
            self.root.after(0, lambda: self.distance_label.config(text="Fehler"))
            logging.error(f"Unerwarteter Fehler bei der Distanzberechnung: {e}")




    # --- 2. Patienten Verwalten Tab (Hinzufügen und Bearbeiten) ---
    def setup_patient_tab(self, tab):
        # --- 1. SEARCH AREA (Top) ---
        search_frame = ttk.LabelFrame(tab, text="Suche & Tools", padding=10)
        search_frame.grid(row=0, column=0, columnspan=3, padx=10, pady=5, sticky='ew')
        
        ttk.Label(search_frame, text="Patient suchen:").pack(side=tk.LEFT, padx=5)
        self.patient_search_entry = ttk.Entry(search_frame, width=25)
        self.patient_search_entry.pack(side=tk.LEFT, padx=5)
        self.patient_search_entry.bind('<Return>', lambda e: self.search_and_load_patient())
        
        ttk.Button(search_frame, text="Laden", command=self.search_and_load_patient).pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="Felder leeren", command=self.reset_patient_form).pack(side=tk.LEFT, padx=5)
        
        # Spacer
        ttk.Label(search_frame, text="   |   ").pack(side=tk.LEFT)
        
        ttk.Button(search_frame, text="♻️ Archiv", command=self.open_archive_manager).pack(side=tk.LEFT, padx=5)
        ttk.Button(search_frame, text="📱 Mobile Scan", command=self.open_scan_dialog).pack(side=tk.LEFT, padx=5)
        
        self.patient_id_to_edit = None

        # --- 2. MAIN CONTENT AREA (Split Left/Right) ---
        
        # Left Side: Inputs
        input_container = ttk.Frame(tab)
        input_container.grid(row=1, column=0, padx=10, pady=5, sticky='n')

        # Personal Data Group
        personal_frame = ttk.LabelFrame(input_container, text="Persönliche Daten", padding=10)
        personal_frame.pack(fill='x', pady=5)

        self.patient_entries = {}
        
        # Helper to create rows
        def add_row(parent, label, field_key, row_idx, width=30):
            ttk.Label(parent, text=f"{label}:").grid(row=row_idx, column=0, padx=5, pady=5, sticky='w')
            entry = ttk.Entry(parent, width=width)
            entry.grid(row=row_idx, column=1, padx=5, pady=5, sticky='ew')
            self.patient_entries[field_key] = entry
            return entry

        add_row(personal_frame, "Anrede", "Anrede", 0)
        add_row(personal_frame, "Vorname", "Vorname", 1)
        add_row(personal_frame, "Nachname", "Nachname", 2)
        add_row(personal_frame, "Versicherungsnummer", "Versicherungsnummer", 3)
        add_row(personal_frame, "Diagnose", "Diagnose", 4)

        # Address Data Group
        address_frame = ttk.LabelFrame(input_container, text="Adresse & Abrechnung", padding=10)
        address_frame.pack(fill='x', pady=5)

        entry_str = add_row(address_frame, "Straße", "Straße", 0)
        entry_str.bind("<FocusOut>", self._start_distance_calculation)
        
        entry_hnr = add_row(address_frame, "Hausnummer", "Hausnummer", 1)
        entry_hnr.bind("<FocusOut>", self._start_distance_calculation)
        
        add_row(address_frame, "Adresszusatz", "Adresszusatz", 2)
        
        # Custom Row for PLZ / Ort
        ttk.Label(address_frame, text="PLZ / Ort:").grid(row=3, column=0, padx=5, pady=5, sticky='w')
        plz_ort_frame = ttk.Frame(address_frame)
        plz_ort_frame.grid(row=3, column=1, sticky='w')
        
        entry_plz = ttk.Entry(plz_ort_frame, width=8)
        entry_plz.pack(side=tk.LEFT, padx=(5, 5))
        self.patient_entries["PLZ"] = entry_plz
        entry_plz.bind("<KeyRelease>", self.autofill_ort)
        entry_plz.bind("<FocusOut>", self._start_distance_calculation)
        
        entry_ort = ttk.Entry(plz_ort_frame, width=20)
        entry_ort.pack(side=tk.LEFT, padx=5)
        self.patient_entries["Ort"] = entry_ort
        entry_ort.bind("<FocusOut>", self._start_distance_calculation)

        # KM Geld
        ttk.Label(address_frame, text="Kilometergeld (€):").grid(row=4, column=0, padx=5, pady=5, sticky='w')
        km_frame = ttk.Frame(address_frame)
        km_frame.grid(row=4, column=1, sticky='w')
        
        entry_km = ttk.Entry(km_frame, width=10)
        entry_km.pack(side=tk.LEFT, padx=(5, 10))
        self.patient_entries["Kilometergeld (€)"] = entry_km
        
        self.distance_label = ttk.Label(km_frame, text="", foreground="blue")
        self.distance_label.pack(side=tk.LEFT)

        # Defaults
        self.patient_entries["Anrede"].insert(0, CONFIG.get('DEFAULT_ANREDE') or "Herr/Frau")
        self.patient_entries["Diagnose"].insert(0, CONFIG.get('DEFAULT_DIAGNOSE') or "Z71")
        self.patient_entries["Kilometergeld (€)"].insert(0, "0.00")

        # Action Buttons Area
        action_frame = ttk.Frame(input_container, padding=10)
        action_frame.pack(fill='x', pady=10)
        
        self.save_patient_button = ttk.Button(action_frame, text="Patient Hinzufügen", command=self.add_patient_gui)
        self.save_patient_button.pack(side=tk.LEFT, fill='x', expand=True, padx=5)
        
        self.delete_patient_button = ttk.Button(action_frame, text="LÖSCHEN", command=self.delete_patient_gui, style='Danger.TButton', state=tk.DISABLED)
        self.delete_patient_button.pack(side=tk.LEFT, padx=5)
        
        ttk.Button(action_frame, text="📂 Ordner", command=self.open_patient_folder).pack(side=tk.LEFT, padx=5)

        # --- 3. SEPARATOR ---
        ttk.Separator(tab, orient='vertical').grid(row=1, column=1, sticky='ns', padx=10, pady=10)

        # --- 4. RIGHT SIDE (Teamup List) ---
        right_frame = ttk.Frame(tab)
        right_frame.grid(row=1, column=2, sticky='nsew', padx=5, pady=5)
        
        # Configure grid weights for resizing
        tab.columnconfigure(2, weight=1)
        tab.rowconfigure(1, weight=1)
        
        ttk.Label(right_frame, text="Neue Patienten (Teamup Check)", font=("Segoe UI", 10, "bold")).pack(pady=(0, 10))
        
        ttk.Button(right_frame, text="Prüfen", command=self.check_new_patients).pack(pady=5, fill='x')
        
        self.new_patients_listbox = tk.Listbox(right_frame, width=40, height=25)
        self.new_patients_listbox.pack(side=tk.LEFT, fill='both', expand=True)
        
        sb = ttk.Scrollbar(right_frame, orient='vertical', command=self.new_patients_listbox.yview)
        sb.pack(side=tk.RIGHT, fill='y')
        self.new_patients_listbox.config(yscrollcommand=sb.set)
        
        self.new_patients_listbox.bind('<Double-1>', self._copy_new_patient_name)
        self.new_patients_listbox.bind('<Button-3>', self._show_new_patients_context_menu)

    def setup_mobile_server(self):
        """Initialisiert und startet den Mobile-Server, falls noch nicht geschehen."""
        if not self.mobile_server:
            self.mobile_server = mobile_connect.MobileServer(self.on_scan_received, self.on_scan_status)
            self.mobile_server.start()

    def on_scan_status(self, message):
        """Callback für Status-Updates vom Server."""
        self.root.after(0, lambda: self._update_scan_status(message))

    def _update_scan_status(self, message):
        if hasattr(self, 'scan_status_label') and self.scan_status_label.winfo_exists():
            self.scan_status_label.config(text=message, foreground="blue")

    def on_scan_received(self, data):
        """Callback vom Server-Thread. Plant GUI-Update im Main-Thread."""
        self.root.after(0, lambda: self._process_scan_data(data))

    def _process_scan_data(self, data):
        """Verarbeitet die empfangenen Daten im GUI-Thread."""
        if not hasattr(self, 'scan_dialog') or not self.scan_dialog.winfo_exists():
            return

        if "error" in data:
            messagebox.showerror("Scan Fehler", data["error"], parent=self.scan_dialog)
            return

        # Mapping definieren
        mapping = {
            "Anrede": "Anrede",
            "Vorname": "Vorname",
            "Nachname": "Nachname",
            "Versicherungsnummer": "Versicherungsnummer",
            "Diagnose": "Diagnose",
            "Straße": "Straße",
            "Hausnummer": "Hausnummer",
            "PLZ": "PLZ",
            "Ort": "Ort"
        }
        
        # Daten in das Hauptformular übertragen
        for json_key, gui_label in mapping.items():
            val = data.get(json_key)
            if val is not None:
                val = str(val).strip()
                if gui_label in self.patient_entries:
                    self.patient_entries[gui_label].delete(0, tk.END)
                    self.patient_entries[gui_label].insert(0, val)
        
        # Trigger distance calculation
        self._start_distance_calculation()
        
        self.set_status("Daten erfolgreich importiert.")
        self.scan_dialog.destroy()

    def open_scan_dialog(self):
        self.setup_mobile_server()
        self.mobile_server.send_ping()
        
        self.scan_dialog = tk.Toplevel(self.root)
        self.scan_dialog.title("Mit Mobile App verbinden")
        self.scan_dialog.geometry("400x500")
        
        # Container for the initial view
        self.scan_initial_frame = ttk.Frame(self.scan_dialog)
        self.scan_initial_frame.pack(fill='both', expand=True)
        
        ttk.Label(self.scan_initial_frame, text="Scannen Sie den Code oder geben Sie die Daten manuell ein:", wraplength=350, justify="center").pack(pady=10)
        
        pil_img = self.mobile_server.get_qr_image().resize((300, 300), Image.Resampling.LANCZOS)
        self.qr_photo = ImageTk.PhotoImage(pil_img, master=self.scan_initial_frame)
        ttk.Label(self.scan_initial_frame, image=self.qr_photo).pack(pady=10)
        
        # NEU: Text-Anzeige für manuelle Eingabe in Flet App
        info_frame = ttk.Frame(self.scan_initial_frame)
        info_frame.pack(pady=5)
        
        ttk.Label(info_frame, text=f"IP:Port:  {self.mobile_server.host_ip}:{self.mobile_server.port}", font=("Consolas", 11, "bold")).pack()
        ttk.Label(info_frame, text=f"Token:    {self.mobile_server.token}", font=("Consolas", 10)).pack()
        
        def on_close_dialog():
            self.mobile_server.set_connection_callback(None)
            self.mobile_server.set_cancel_callback(None)
            self.scan_dialog.destroy()
            
        self.scan_dialog.protocol("WM_DELETE_WINDOW", on_close_dialog)

        def trigger_scan_action():
            if not self.scan_initial_frame.winfo_exists(): return
            self.mobile_server.trigger_scan()
            
            # Switch to loading view
            self.scan_initial_frame.destroy()
            
            loading_frame = ttk.Frame(self.scan_dialog)
            loading_frame.pack(fill='both', expand=True, padx=20, pady=20)
            
            ttk.Label(loading_frame, text="Waiting for response...", font=("Segoe UI", 16)).pack(pady=(80, 20))
            
            pb = ttk.Progressbar(loading_frame, mode='indeterminate', length=200)
            pb.pack(pady=10)
            pb.start(15)
            
            ttk.Button(loading_frame, text="Cancel", command=on_close_dialog).pack(pady=40)

        def on_mobile_connected():
            logging.info("Pong received. Switching to loading UI.")
            self.root.after(0, perform_auto_scan)

        def on_mobile_cancelled():
            self.root.after(0, lambda: self.set_status("Vorgang am mobilen Gerät abgebrochen."))
            self.root.after(0, on_close_dialog)

        def perform_auto_scan():
            if hasattr(self, 'scan_dialog') and self.scan_dialog.winfo_exists():
                if hasattr(self, 'scan_initial_frame') and self.scan_initial_frame.winfo_exists():
                    self.set_status("Mobile App verbunden. Scan wird angefordert...")
                    trigger_scan_action()

        self.mobile_server.set_connection_callback(on_mobile_connected)
        self.mobile_server.set_cancel_callback(on_mobile_cancelled)

        ttk.Button(self.scan_initial_frame, text="Request", command=trigger_scan_action).pack(pady=5)
        
        self.scan_status_label = ttk.Label(self.scan_initial_frame, text="Warte auf Scan...", font=("Segoe UI", 10, "italic"))
        self.scan_status_label.pack(pady=10)
        ttk.Button(self.scan_initial_frame, text="Abbrechen", command=on_close_dialog).pack(pady=10)

    def open_archive_manager(self):
        """Öffnet ein Fenster zum Suchen und Reaktivieren von archivierten Patienten."""
        archive_win = tk.Toplevel(self.root)
        archive_win.title("Archivierte Patienten")
        archive_win.geometry("600x500")
        
        # --- Search Area ---
        search_frame = ttk.Frame(archive_win)
        search_frame.pack(fill='x', padx=10, pady=10)
        
        ttk.Label(search_frame, text="Suche:").pack(side=tk.LEFT, padx=(0, 5))
        search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=search_var)
        search_entry.pack(side=tk.LEFT, fill='x', expand=True, padx=(0, 5))
        search_entry.focus_set()
        
        ttk.Label(archive_win, text="Archivierte Patienten (Doppelklick zum Reaktivieren):", font=("Segoe UI", 10, "bold")).pack(pady=(0, 5))
        
        list_frame = ttk.Frame(archive_win)
        list_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        lb = tk.Listbox(list_frame, width=80, height=20)
        lb.pack(side="left", fill="both", expand=True)
        
        scroll = ttk.Scrollbar(list_frame, command=lb.yview)
        scroll.pack(side="right", fill="y")
        lb.config(yscrollcommand=scroll.set)
        
        # Daten laden
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute("SELECT id, vorname, nachname FROM patienten WHERE is_archived = 1 ORDER BY nachname")
        all_archived_patients = cursor.fetchall()
        conn.close()
        
        def update_list(*args):
            search_term = search_var.get().lower()
            lb.delete(0, tk.END)
            for pid, vn, nn in all_archived_patients:
                full_str = f"{nn} {vn} (ID: {pid})"
                if search_term in full_str.lower():
                    lb.insert(tk.END, full_str)

        search_var.trace_add("write", update_list)
        update_list() # Initial fill
            
        def reactivate():
            sel = lb.curselection()
            if not sel: return
            
            item_text = lb.get(sel[0])
            try:
                pid_str = item_text.rsplit("(ID: ", 1)[1].replace(")", "")
                pid = int(pid_str)
            except (IndexError, ValueError):
                return

            patient_data = next((p for p in all_archived_patients if p[0] == pid), None)
            if not patient_data: return
            
            _, vn, nn = patient_data
            folder_name = f"{nn} {vn}"
            
            if messagebox.askyesno("Reaktivieren", f"Möchten Sie '{folder_name}' reaktivieren?\nDer Ordner wird aus dem Archiv zurückverschoben."):
                # 1. DB Update
                conn = sqlite3.connect(DATABASE_NAME)
                cursor = conn.cursor()
                cursor.execute("UPDATE patienten SET is_archived = 0 WHERE id = ?", (pid,))
                
                # NEU: Dazugehörige Stammdatenleistung ebenfalls reaktivieren
                cursor.execute("UPDATE stammdaten_leistungen SET is_archived = 0 WHERE kurzname=?", (nn,))
                
                conn.commit()
                conn.close()
                
                # 2. Ordner verschieben
                archive_dir = CONFIG.get('ARCHIVE_DIR')
                base_dir = CONFIG.get('PATIENT_BASE_DIR')
                src = os.path.join(archive_dir, folder_name)
                dst = os.path.join(base_dir, folder_name)
                
                if os.path.exists(src):
                    shutil.move(src, dst)
                    messagebox.showinfo("Erfolg", "Patient und Ordner erfolgreich reaktiviert.")
                else:
                    messagebox.showwarning("Hinweis", "Patient wurde in der Datenbank reaktiviert, aber der Ordner wurde im Archiv nicht gefunden.")
                
                archive_win.destroy()
                self.search_patients() # Hauptliste aktualisieren
                self.refresh_all_stammdaten_ui() # NEU: Stammdaten-UI aktualisieren

        lb.bind("<Double-1>", lambda e: reactivate())
        ttk.Button(archive_win, text="Ausgewählten Patienten Reaktivieren", command=reactivate).pack(pady=10)

    def check_new_patients(self):
        """
        Holt Teamup-Events des gewählten Monats und vergleicht sie mit der DB.
        Events, deren Titel keinen bekannten Nachnamen enthalten, werden gelistet.
        """
        mode = CONFIG.get('AUTO_DATE_SELECTOR', 'Auto')
        
        if mode == 'Manual':
            start_date = CONFIG.get('MANUAL_DATE_START', '')
            end_date = CONFIG.get('MANUAL_DATE_END', '')
            if not start_date or not end_date:
                messagebox.showerror("Fehler", "Manueller Datumsbereich ist nicht konfiguriert (siehe Einstellungen).")
                return
        else:
            try:
                month = int(self.selected_invoice_month.get())
                year = int(self.selected_invoice_year.get())
                
                start_date = f"{year}-{month:02d}-01"
                last_day = calendar.monthrange(year, month)[1]
                end_date = f"{year}-{month:02d}-{last_day}"
                
            except ValueError:
                messagebox.showerror("Fehler", "Ungültiges Datum ausgewählt.")
                return

        self.set_status("Lade Teamup-Daten und vergleiche mit Datenbank...", 0)
        self.root.update_idletasks()

        events = search_teamup_events("", start_date=start_date, end_date=end_date)
        
        if not events:
            self.new_patients_listbox.delete(0, tk.END)
            self.new_patients_listbox.insert(tk.END, "Keine Termine gefunden.")
            self.set_status("Keine Termine im Zeitraum gefunden.")
            return

        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute("SELECT nachname FROM patienten WHERE is_archived = 0 OR is_archived IS NULL")
        db_lastnames = [r[0].strip().lower() for r in cursor.fetchall() if r[0] and r[0].strip()]
        
        conn.close()

        blacklist_terms = [term.lower() for term in self.session_blacklist]
        unknown_patients = []
        
        for title, date_str, _, _ in events:
            title_lower = title.lower()
            
            if any(term in title_lower for term in blacklist_terms):
                continue

            is_known = False
            
            for db_name in db_lastnames:
                if db_name in title_lower:
                    is_known = True
                    break
            
            if not is_known:
                unknown_patients.append(f"{date_str}: {title}")

        self.new_patients_listbox.delete(0, tk.END)
        if unknown_patients:
            for entry in unknown_patients:
                self.new_patients_listbox.insert(tk.END, entry)
            self.set_status(f"{len(unknown_patients)} potenzielle neue Patienten gefunden.")
        else:
            self.new_patients_listbox.insert(tk.END, "Alle Patienten bekannt.")
            self.set_status("Alle Termine konnten zugeordnet werden.")

    def _copy_new_patient_name(self, event):
        """Kopiert den Namen aus der Liste in das Nachname-Feld."""
        selection = self.new_patients_listbox.curselection()
        if selection:
            text = self.new_patients_listbox.get(selection[0])
            if ": " in text:
                name_part = text.split(": ", 1)[1]
            else:
                name_part = text
            
            self.patient_entries["Nachname"].delete(0, tk.END)
            self.patient_entries["Nachname"].insert(0, name_part)

    def _show_new_patients_context_menu(self, event):
        """Zeigt Kontextmenü für die Liste neuer Patienten."""
        try:
            index = self.new_patients_listbox.nearest(event.y)
            if index == -1: return
            
            self.new_patients_listbox.selection_clear(0, tk.END)
            self.new_patients_listbox.selection_set(index)
            self.new_patients_listbox.activate(index)
            
            menu = tk.Menu(self.root, tearoff=0)
            menu.add_command(label="Kopieren", command=lambda: self._copy_new_patient_to_clipboard(index))
            menu.add_command(label="Zu Blacklist hinzufügen", command=lambda: self._blacklist_new_patient(index))
            menu.tk_popup(event.x_root, event.y_root)
        except Exception:
            pass

    def _copy_new_patient_to_clipboard(self, index):
        text = self.new_patients_listbox.get(index)
        name_part = text.split(": ", 1)[1] if ": " in text else text
        self.root.clipboard_clear()
        self.root.clipboard_append(name_part)
        self.set_status("Name in Zwischenablage kopiert.")

    def _blacklist_new_patient(self, index):
        text = self.new_patients_listbox.get(index)
        name_part = text.split(": ", 1)[1] if ": " in text else text
            
        term = simpledialog.askstring("Blacklist", "Begriff für Blacklist für diese Sitzung eingeben (z.B. 'Meeting', 'Urlaub'):", initialvalue=name_part, parent=self.root)
        
        if term and term.strip():
            clean_term = term.strip()
            self.session_blacklist.add(clean_term)
            self.set_status(f"'{clean_term}' zur Blacklist für diese Sitzung hinzugefügt.")
            # Refresh the list to reflect the change
            self.check_new_patients()

    def search_and_load_patient(self):
        # ... (Funktion bleibt unverändert)
        search_term = self.patient_search_entry.get().strip()
        self.patient_search_entry.delete(0, tk.END)

        if not search_term:
            messagebox.showwarning("Suche", "Bitte geben Sie einen Suchbegriff ein.")
            return
            
        results = get_patient_data(search_term) 
        
        if not results:
            self.reset_patient_form()
            messagebox.showinfo("Suche", f"Kein Patient mit '{search_term}' gefunden.")
            return
            
        if len(results) > 1:
            # Mehrere Ergebnisse, zeige Auswahlfenster
            self._show_patient_selection_dialog(results)
        else:
            # Ein klares Ergebnis
            patient_data_tuple = results[0]
            self.load_patient_data_to_form(patient_data_tuple[0], patient_data_tuple)


    def _show_patient_selection_dialog(self, results):
        # ... (Funktion bleibt unverändert)
        selection_window = tk.Toplevel(self.root)
        selection_window.title("Patienten Auswahl")
        
        ttk.Label(selection_window, text="Mehrere Patienten gefunden. Bitte wählen Sie einen aus:").pack(pady=10)
        
        listbox = tk.Listbox(selection_window, width=70, height=10)
        listbox.pack(padx=10, pady=5)
        
        for patient_data_tuple in results:
            display_text = f"ID {patient_data_tuple[0]} - {patient_data_tuple[2]} {patient_data_tuple[1]} ({patient_data_tuple[7]})"
            listbox.insert(tk.END, display_text)
            
        def on_select():
            selection = listbox.curselection()
            if selection:
                index = selection[0]
                patient_id = results[index][0]
                patient_data = results[index]
                self.load_patient_data_to_form(patient_id, patient_data)
                selection_window.destroy()
            else:
                messagebox.showwarning("Auswahl", "Bitte wählen Sie einen Patienten.")

        ttk.Button(selection_window, text="Auswählen und Laden", command=on_select).pack(pady=10)

        # Zentrieren des Fensters
        selection_window.update_idletasks()
        width = selection_window.winfo_width()
        height = selection_window.winfo_height()
        x = (selection_window.winfo_screenwidth() // 2) - (width // 2)
        y = (selection_window.winfo_screenheight() // 2) - (height // 2)
        selection_window.geometry('{}x{}+{}+{}'.format(width, height, x, y))


    def load_patient_data_to_form(self, patient_id, data):
        # ID, vorname, nachname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, last_selected_kurznamen
        fields = [
            "Anrede", "Vorname", "Nachname", "Versicherungsnummer", 
            "Straße", "Hausnummer", "Adresszusatz", "PLZ", "Ort", 
            "Diagnose", "Kilometergeld (€)"
        ]
        
        for i, field in enumerate(fields):
            self.patient_entries[field].delete(0, tk.END)
            
            # Kilometergeld (Index 11) muss mit Komma formatiert werden
            if field == "Kilometergeld (€)":
                value = f"{data[11]:.2f}".replace('.', ',') if data[11] is not None else "0,00"
            # Adresszusatz (Index 5) 
            elif field == "Adresszusatz":
                value = data[5] if data[5] is not None else ""
            # Alle anderen Felder
            elif i < 3: # Anrede(8), Vorname(1), Nachname(2)
                 value = data[i+5] if i == 0 else data[i] 
            else: # Restliche Felder
                 # Anrede ist Index 8, Vorname 1, Nachname 2, VersNr 9, Str 3, HNr 4, AdrZu 5, PLZ 6, Ort 7, Dia 10, KM 11
                 db_index = [8, 1, 2, 9, 3, 4, 5, 6, 7, 10, 11] # Korrekte Mapping
                 value = data[db_index[i]]
                 if field == "Adresszusatz" and value is None: value = "" # Adresszusatz kann None sein

            # Korrigiertes Mapping
            mapping = {
                "Anrede": data[8], "Vorname": data[1], "Nachname": data[2], 
                "Versicherungsnummer": data[9], "Straße": data[3], "Hausnummer": data[4], 
                "Adresszusatz": data[5] if data[5] is not None else "", "PLZ": data[6], 
                "Ort": data[7], "Diagnose": data[10], 
                "Kilometergeld (€)": f"{data[11]:.2f}".replace('.', ',') if data[11] is not None else "0,00"
            }
            
            self.patient_entries[field].insert(0, mapping[field])

        # Button-Text und ID setzen
        self.save_patient_button.config(text="Patient Aktualisieren", command=self.update_patient_gui)
        self.patient_id_to_edit = patient_id
        
        # NEU: Lösch-Button aktivieren
        self.delete_patient_button.config(state=tk.NORMAL) 
        
        messagebox.showinfo("Geladen", f"Patient ID {patient_id} ({data[2]} {data[1]}) erfolgreich zur Bearbeitung geladen.")

    def open_patient_folder(self):
        """Öffnet den Ordner des aktuell geladenen Patienten im Explorer/Finder."""
        if not self.patient_id_to_edit:
            messagebox.showwarning("Info", "Bitte laden Sie zuerst einen Patienten.")
            return
        
        nachname = self.patient_entries["Nachname"].get().strip()
        vorname = self.patient_entries["Vorname"].get().strip()
        folder_name = f"{nachname} {vorname}"
        output_folder = os.path.expanduser(CONFIG.get('PATIENT_BASE_DIR'))
        path = os.path.join(output_folder, folder_name)
        
        if os.path.exists(path):
            if sys.platform == 'win32': os.startfile(path)
            else: subprocess.Popen(['xdg-open', path])
        else:
            messagebox.showinfo("Info", f"Noch kein Ordner für '{folder_name}' vorhanden.\n(Wird erst bei der ersten Rechnung erstellt)")

    def reset_patient_form(self):
        fields = [
            "Anrede", "Vorname", "Nachname", "Versicherungsnummer", 
            "Straße", "Hausnummer", "Adresszusatz", "PLZ", "Ort", 
            "Diagnose", "Kilometergeld (€)"
        ]
        for field in fields:
            self.patient_entries[field].delete(0, tk.END)
            
        self.patient_entries["Anrede"].insert(0, CONFIG.get('DEFAULT_ANREDE') or "Herr/Frau")
        self.patient_entries["Diagnose"].insert(0, CONFIG.get('DEFAULT_DIAGNOSE') or "Z71")
        self.patient_entries["Kilometergeld (€)"].insert(0, "0.00")
        
        # Button-Text und ID zurücksetzen
        self.save_patient_button.config(text="Patient Hinzufügen", command=self.add_patient_gui)
        self.patient_id_to_edit = None
        
        # NEU: Lösch-Button deaktivieren
        self.delete_patient_button.config(state=tk.DISABLED)
        
        messagebox.showinfo("Zurückgesetzt", "Formular wurde zurückgesetzt.")

    def update_patient_gui(self):
        """Wrapper für add_patient_gui, um Update-Aktionen zu starten."""
        self.add_patient_gui(is_update=True)


    def add_patient_gui(self, is_update=False):
        """
        Fügt einen neuen Patienten hinzu oder aktualisiert einen bestehenden.
        Die Logik für Hinzufügen und Aktualisieren wurde in dieser Funktion konsolidiert.
        """
        data = {k: v.get().strip() for k, v in self.patient_entries.items()}
        
        try:
            # Pflichtfelder
            nachname = data.get("Nachname")
            vorname = data.get("Vorname")
            plz = data.get("PLZ")
            
            # Sicherstellen, dass die Dezimaltrennzeichen korrekt sind und Leere in 0.0 umgewandelt werden
            km_string = data.get("Kilometergeld (€)").replace(',', '.') if data.get("Kilometergeld (€)") else '0.0'
            kilometergeld = float(km_string)
            
            # Basis-Validierung
            if not nachname or not vorname or not plz:
                messagebox.showwarning("Achtung", "Nachname, Vorname und PLZ sind Pflichtfelder.")
                return

            # Alle anderen Felder
            strasse = data.get("Straße")
            hausnummer = data.get("Hausnummer")
            adresszusatz = data.get("Adresszusatz")
            ort = data.get("Ort")
            anrede = data.get("Anrede")
            versicherungsnummer = data.get("Versicherungsnummer")
            diagnose = data.get("Diagnose")

        except ValueError:
            messagebox.showerror("Fehler", "Kilometergeld muss eine gültige Zahl sein.")
            return

        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        try:
            # --- AKTUALISIEREN (UPDATE) ---
            if self.patient_id_to_edit:
                patient_id = self.patient_id_to_edit
                
                # NEU: Hole die aktuellen einzigartigen Felder des Patienten aus der DB
                cursor.execute("""
                SELECT vorname, nachname, plz FROM patienten WHERE id = ?
                """, (patient_id,))
                original_data = cursor.fetchone()

                # Prüfen, ob die kritischen Felder geändert wurden
                vorname_changed = (original_data[0] != vorname)
                nachname_changed = (original_data[1] != nachname)
                plz_changed = (original_data[2] != plz)
                
                # FIX: Führe die Kollisionsprüfung NUR aus, wenn sich VORNAME, NACHNAME oder PLZ geändert haben.
                # Dadurch können nicht-kritische Felder wie Adresszusatz oder Kilometergeld 
                # aktualisiert werden, selbst wenn ein (vermutlich versehentlich angelegter) 
                # Duplikat-Patient existiert.
                if vorname_changed or nachname_changed or plz_changed:
                    # 2. PRÜFUNG: Kollidieren die NEUEN Daten mit einem ANDEREN Patienten?
                    cursor.execute(""" 
                    SELECT id FROM patienten WHERE nachname = ? AND vorname = ? AND plz = ? AND id != ? 
                    """, (nachname, vorname, plz, patient_id))
                    
                    if cursor.fetchone(): 
                        # Konflikt mit einem ANDEREN Patienten gefunden
                        messagebox.showerror( 
                        "Fehler: Datenkonflikt", 
                        f"Die aktualisierten Daten ('{vorname} {nachname}' mit PLZ {plz}) kollidieren mit einem ANDEREN, bereits existierenden Patienten. Update abgebrochen." 
                        ) 
                        return # Funktion hier beenden
                    # Wenn die Schlüsseldaten geändert wurden und kein Konflikt besteht, fahre mit dem Update fort
                
                # 3. Führe das Update aus
                cursor.execute("""
                UPDATE patienten 
                SET nachname=?, vorname=?, strasse=?, hausnummer=?, adresszusatz=?, plz=?, ort=?, anrede=?, versicherungsnummer=?, diagnose=?, kilometergeld=?
                WHERE id=?
                """, (nachname, vorname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, patient_id))
                conn.commit()
                messagebox.showinfo("Erfolg", f"Patient '{vorname} {nachname}' (ID: {patient_id}) erfolgreich aktualisiert.")
                self.reset_patient_form() 
                self.search_patients() # Aktualisiere Hauptliste

            # --- HINZUFÜGEN (INSERT) ---
            else: 
                # 1. PRÜFUNG: Existiert der Patient bereits?
                cursor.execute("""
                SELECT id FROM patienten WHERE nachname = ? AND vorname = ? AND plz = ?
                """, (nachname, vorname, plz))
                
                if cursor.fetchone():
                    # Patient existiert bereits
                    raise sqlite3.IntegrityError("Patient exists") 
                    
                # 2. Führe das INSERT aus
                cursor.execute("""
                INSERT INTO patienten (vorname, nachname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, last_selected_kurznamen, invoiced_since_reset) 
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (vorname, nachname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, '', 0))
                conn.commit()
                messagebox.showinfo("Erfolg", f"Patient '{vorname} {nachname}' erfolgreich hinzugefügt.")
                self.reset_patient_form() 
                self.search_patients() # Aktualisiere Hauptliste
                
        except sqlite3.IntegrityError:
             # Fängt den Fehler des INSERT-Falls ab
            messagebox.showerror("Fehler", f"Patient \"{vorname} {nachname}\" mit dieser PLZ existiert bereits in der Datenbank.")
        except Exception as e:
            messagebox.showerror("Fehler", f"Unbekannter Datenbankfehler: {e}")
        finally:
            conn.close()

    
    def _delete_patient_from_db(self, patient_id):
        """Löscht den Patienten und seine Leistungen aus der Datenbank."""
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        try:
            # 1. Alle Leistungen löschen (WICHTIG für Datenintegrität)
            self._delete_all_patient_leistungen(patient_id) 
            
            # 2. Patienten löschen
            cursor.execute("DELETE FROM patienten WHERE id = ?", (patient_id,))
            conn.commit()
            return True
        except Exception as e:
            logging.error(f"FEHLER beim Löschen von Patient ID {patient_id}: {e}", exc_info=True)
            messagebox.showerror("Fehler", f"Fehler beim Löschen des Patienten: {e}")
            return False
        finally:
            conn.close()

    def delete_patient_gui(self):
        """Bestätigt und löscht den aktuell geladenen Patienten."""
        if not self.patient_id_to_edit:
            messagebox.showwarning("Achtung", "Kein Patient zum Löschen ausgewählt.")
            return

        nachname = self.patient_entries["Nachname"].get().strip()
        vorname = self.patient_entries["Vorname"].get().strip()
        patient_id = self.patient_id_to_edit

        if messagebox.askyesno("Bestätigung Löschen", 
                               f"Sind Sie sicher, dass Sie den Patienten '{vorname} {nachname}' (ID: {patient_id}) und ALLE seine Leistungen dauerhaft löschen möchten?\n\nDIESER SCHRITT KANN NICHT RÜCKGÄNGIG GEMACHT WERDEN!"):
            
            if self._delete_patient_from_db(patient_id):
                messagebox.showinfo("Erfolg", f"Patient '{vorname} {nachname}' wurde erfolgreich gelöscht.")
                self.reset_patient_form()
                # Aktualisiere die Haupt-Suchergebnisse (Generieren-Tab)
                self.search_patients()
    
    
    # --- HILFSFUNKTIONEN FÜR LEISTUNGEN --- 
    def _delete_all_patient_leistungen(self, patient_id):
        """Löscht ALLE Leistungen des angegebenen Patienten ohne GUI-Interaktion."""
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        try:
            cursor.execute("DELETE FROM leistungen WHERE patient_id = ?", (patient_id,))
            conn.commit()
            return True
        except Exception as e:
            logging.error(f"FEHLER: Fehler beim Löschen aller Leistungen für Patient {patient_id}: {e}")
            return False
        finally:
            conn.close()


    def _get_leistung_insertion_params(self):
        """Hilfsfunktion zur Vorbereitung von Betrag und KM-Geld."""
        # KM-Geld basierend auf Checkbox
        if self.use_km_money_var.get():
            km_geld = self.get_current_kilometergeld()
        else:
            km_geld = 0.0

        # Wenn der Editor nicht sichtbar ist, ignorieren wir manuelle Eingaben
        if not hasattr(self, 'editor_window') or not self.editor_window or not self.editor_window.winfo_exists():
            return 0.0, False, km_geld

        manual_betrag_str = self.amount_entry.get().strip().replace(',', '.')
        try:
            manual_betrag = float(manual_betrag_str)
            use_manual_override = manual_betrag > 0.009
        except ValueError:
            use_manual_override = False
            manual_betrag = 0.0
        return manual_betrag, use_manual_override, km_geld

    def _reset_leistung_selection(self):
        """Hilfsfunktion zum Zurücksetzen der Button-Auswahl in der GUI."""
        self.selected_leistungs_kurznamen.clear()
        for widget in self.leistung_button_frame.winfo_children():
            # Reset unconditionally to prevent "stuck" selections
            if hasattr(widget, 'kurzname'):
                widget.config(style='TButton')
                widget.is_selected = False

    def add_leistung_to_db(self, patient_id, datum_str, time_from, time_to, kurzname, standard_betrag, manual_betrag, use_manual_override, km_geld, custom_description=None):
        """Fügt eine einzelne Leistung in die DB ein."""
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        try:
            datum_db = datetime.datetime.strptime(datum_str, '%d.%m.%Y').strftime('%Y-%m-%d')

            betrag = manual_betrag if use_manual_override else standard_betrag
            end_betrag = betrag + km_geld 

            # Beschreibung basierend auf Textbox, manuellem Override oder Stammdaten
            if custom_description:
                beschreibung = custom_description
            elif use_manual_override:
                beschreibung = f"Manuelle Eingabe"
            else:
                beschreibung = self.stammdaten_betraege.get(kurzname, kurzname)

            cursor.execute("""
            INSERT INTO leistungen (patient_id, datum, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag)
            VALUES (?, ?, ?, ?, ?, ?)
            """, (patient_id, datum_db, time_from, time_to, beschreibung, end_betrag))
            conn.commit()
            return True
        except Exception as e:
            logging.error(f"FEHLER: Fehler beim Speichern der Leistung {kurzname}: {e}")
            return False
        finally:
            conn.close()

    def _insert_multiple_leistungen(self, patient_id, events_list):
        """Fügt mehrere Leistungen aus Teamup-Einträgen in die DB ein."""
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        total_success_count = 0
        if self.use_km_money_var.get():
            km_geld = self.get_current_kilometergeld()
        else:
            km_geld = 0.0

        for title, date_str, time_from, time_to in events_list:
            
            # WICHTIG: Die ausgewählten Leistungen (self.selected_leistungs_kurznamen) werden auf jeden Termin angewendet.
            for kurzname in self.selected_leistungs_kurznamen:
                
                # Hole den Betrag für diese Leistung
                # Der Kurzname in stammdaten_betraege ist im Format 'Kurzname - Beschreibung'
                stammdaten_key = [k for k in self.stammdaten_betraege.keys() if k.startswith(kurzname + ' -')]
                if not stammdaten_key:
                    logging.warning(f"Stammdaten für '{kurzname}' nicht gefunden. Überspringe.")
                    continue
                
                standard_betrag = self.stammdaten_betraege[stammdaten_key[0]]
                end_betrag = standard_betrag + km_geld
                
                # Die Beschreibung der Leistung wird aus den Stammdaten übernommen
                final_beschreibung = stammdaten_key[0] 
                
                try:
                    datum_db = datetime.datetime.strptime(date_str, '%d.%m.%Y').strftime('%Y-%m-%d')

                    # Die Beschreibung wird aus den Stammdaten übernommen, nicht aus dem Teamup-Titel
                    cursor.execute("""
                    INSERT INTO leistungen (patient_id, datum, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag)
                    VALUES (?, ?, ?, ?, ?, ?)
                    """, (patient_id, datum_db, time_from, time_to, final_beschreibung, end_betrag))
                    total_success_count += 1
                except Exception as e:
                    logging.error(f"Fehler beim Speichern des Termins {date_str} für Leistung {kurzname}: {e}")

        conn.commit()
        conn.close()
        return total_success_count 
    
    # --- ENDE HILFSFUNKTIONEN FÜR LEISTUNGEN --- 

    # --- 3. Leistungen Hinzufügen/Prüfen Tab (Mit Uhrzeiten und Teamup) --- 
    # ... (Rest der Klasse bleibt unverändert)

    def _on_canvas_configure(self, event):
        """Passt die Breite des inneren Frames an die Breite des Canvas an."""
        self.leistung_canvas.itemconfig(self.window_id, width=event.width)

    def setup_leistung_tab(self, tab):
        # Grid configuration
        tab.columnconfigure(0, weight=1)
        tab.rowconfigure(3, weight=1) # Treeview expands (Row 3)

        # --- 1. Patient Info ---
        info_frame = ttk.Frame(tab)
        info_frame.grid(row=0, column=0, sticky='ew', padx=10, pady=5)
        ttk.Label(info_frame, text="Aktueller Patient:", font=('Segoe UI', 10)).pack(side=tk.LEFT)
        self.leistung_patient_label = ttk.Label(info_frame, text="Bitte Patient in Tab 1 auswählen", foreground='red', font=('Segoe UI', 10, 'bold'))
        self.leistung_patient_label.pack(side=tk.LEFT, padx=10)

        # --- 2. Main Actions (Teamup & Finish) ---
        main_action_frame = ttk.Frame(tab)
        main_action_frame.grid(row=1, column=0, sticky='ew', padx=10, pady=10)
        
        # Styles
        self.ttk_style.configure('Prominent.TButton', font=('Segoe UI', 12, 'bold'))
        self.ttk_style.configure('Finish.TButton', font=('Segoe UI', 12, 'bold'), foreground='green')

        self.teamup_button = ttk.Button(main_action_frame, text="📅 Teamup-Termine Importieren", command=self.open_teamup_search, style='Prominent.TButton')
        self.teamup_button.pack(side=tk.LEFT, fill='x', expand=True, padx=(0, 10), ipady=5)

        finish_btn = ttk.Button(main_action_frame, text="✅ Fertig / Weiter zum Drucken", command=self._switch_to_generate_tab, style='Finish.TButton')
        finish_btn.pack(side=tk.LEFT, fill='x', expand=True, padx=(10, 0), ipady=5)

        # --- 3. Leistungsauswahl (Stammdaten Buttons) ---
        selection_frame = ttk.LabelFrame(tab, text="Leistung auswählen")
        selection_frame.grid(row=2, column=0, sticky='ew', padx=10, pady=5)
        
        self.leistung_canvas = tk.Canvas(selection_frame, height=200)
        self.leistung_canvas.pack(side="left", fill="x", expand=True, padx=5, pady=5)

        ls_scrollbar = ttk.Scrollbar(selection_frame, orient="vertical", command=self.leistung_canvas.yview)
        ls_scrollbar.pack(side="right", fill="y", pady=5)

        self.leistung_canvas.configure(yscrollcommand=ls_scrollbar.set)
        self.leistung_canvas.bind('<Configure>', self._on_canvas_configure)

        self.leistung_button_frame = ttk.Frame(self.leistung_canvas)
        self.window_id = self.leistung_canvas.create_window((0, 0), window=self.leistung_button_frame, anchor="nw")

        # --- 4. Treeview (Existing Services) ---
        tree_frame = ttk.LabelFrame(tab, text="Erfasste Leistungen")
        tree_frame.grid(row=3, column=0, sticky='nsew', padx=10, pady=5)
        
        columns = ('ID', 'Datum', 'Von', 'Bis', 'Beschreibung', 'Betrag')
        self.leistung_tree = ttk.Treeview(tree_frame, columns=columns, show='headings', selectmode='browse')
        
        tree_scroll = ttk.Scrollbar(tree_frame, orient="vertical", command=self.leistung_tree.yview)
        self.leistung_tree.configure(yscrollcommand=tree_scroll.set)
        
        self.leistung_tree.pack(side=tk.LEFT, fill='both', expand=True)
        tree_scroll.pack(side=tk.RIGHT, fill='y')

        for col in columns:
            self.leistung_tree.heading(col, text=col)
        
        self.leistung_tree.column('ID', width=40, anchor='center')
        self.leistung_tree.column('Datum', width=100, anchor='center')
        self.leistung_tree.column('Von', width=60, anchor='center')
        self.leistung_tree.column('Bis', width=60, anchor='center')
        self.leistung_tree.column('Beschreibung', width=300, anchor='w')
        self.leistung_tree.column('Betrag', width=120, anchor='e')
        
        self.leistung_tree.bind('<<TreeviewSelect>>', self.select_leistung_for_edit)

        # --- 5. Summary & List Controls ---
        summary_frame = ttk.Frame(tab)
        summary_frame.grid(row=4, column=0, sticky='ew', padx=10, pady=5)

        self.summary_label = ttk.Label(summary_frame, text="Gesamtsumme: €0.00 | Nicht abgerechnete Leistungen: 0", font=('Segoe UI', 10, 'bold'))
        self.summary_label.pack(side=tk.LEFT)

        ttk.Button(summary_frame, text="Alle Löschen", command=self.delete_all_leistungen_gui).pack(side=tk.RIGHT, padx=5)
        ttk.Button(summary_frame, text="Löschen", command=self.delete_leistung_gui).pack(side=tk.RIGHT, padx=5)
        ttk.Button(summary_frame, text="Bearbeiten", command=self.load_leistung_for_edit).pack(side=tk.RIGHT, padx=5)

        # --- 6. Date/Time & Add & Manual ---
        bottom_group = ttk.LabelFrame(tab, text="Hinzufügen / Bearbeiten")
        bottom_group.grid(row=5, column=0, sticky='ew', padx=10, pady=10)

        # Date/Time/KM
        dt_frame = ttk.Frame(bottom_group)
        dt_frame.pack(fill='x', padx=5, pady=5)
        
        ttk.Label(dt_frame, text="Datum:").pack(side=tk.LEFT)
        self.date_entry = ttk.Entry(dt_frame, width=12)
        self.date_entry.pack(side=tk.LEFT, padx=5)
        self.date_entry.insert(0, datetime.date.today().strftime("%d.%m.%Y"))

        ttk.Label(dt_frame, text="Von:").pack(side=tk.LEFT)
        self.time_from_entry = ttk.Entry(dt_frame, width=6)
        self.time_from_entry.pack(side=tk.LEFT, padx=5)
        self.time_from_entry.insert(0, "11:00")

        ttk.Label(dt_frame, text="Bis:").pack(side=tk.LEFT)
        self.time_to_entry = ttk.Entry(dt_frame, width=6)
        self.time_to_entry.pack(side=tk.LEFT, padx=5)
        self.time_to_entry.insert(0, "11:50")

        self.km_check = ttk.Checkbutton(dt_frame, text="inkl. KM-Geld", variable=self.use_km_money_var)
        self.km_check.pack(side=tk.LEFT, padx=15)

        # Buttons
        btn_frame = ttk.Frame(bottom_group)
        btn_frame.pack(fill='x', padx=5, pady=5)

        self.add_selection_button = ttk.Button(btn_frame, text="Ausgewählte Leistung(en) hinzufügen", command=self.add_leistung_gui)
        self.add_selection_button.pack(side=tk.LEFT, padx=10)
        
        ttk.Button(btn_frame, text="Auswahl Reset", command=lambda: self._reset_leistung_selection()).pack(side=tk.LEFT, padx=5)

        self.toggle_editor_btn = ttk.Button(btn_frame, text="Manuelle Eingabe / Editor öffnen", command=self.open_manual_editor)
        self.toggle_editor_btn.pack(side=tk.RIGHT, padx=10)

        # --- Hidden Widgets for Compatibility ---
        # These are needed because add_leistung_gui references them even if the editor window isn't open
        self.editor_frame = ttk.Frame(tab) # Dummy frame
        self.description_text = tk.Text(self.editor_frame, height=1, width=1)
        self.amount_entry = ttk.Entry(self.editor_frame)
        self.amount_entry.insert(0, "0.00")
        
        # Backups for restore when Toplevel closes
        self.main_description_text = self.description_text
        self.main_amount_entry = self.amount_entry
        self.main_add_leistung_button = self.add_selection_button
        self.add_leistung_button = self.add_selection_button # Initial reference

    def open_manual_editor(self):
        if hasattr(self, 'editor_window') and self.editor_window and self.editor_window.winfo_exists():
            self.editor_window.lift()
            self.editor_window.focus_force()
            return

        self.editor_window = tk.Toplevel(self.root)
        self.editor_window.title("Manuelle Leistungseingabe / Editor")
        self.editor_window.geometry("700x450")
        self.editor_window.attributes("-topmost", True)
        self.editor_window.protocol("WM_DELETE_WINDOW", self._close_manual_editor_window)
        
        # 3.1 Loader (Stammdaten in Editor laden)
        loader_frame = ttk.Frame(self.editor_window)
        loader_frame.pack(fill='x', pady=10, padx=10)
        ttk.Label(loader_frame, text="Vorlage aus Stammdaten laden:").pack(side=tk.LEFT)
        self.stammdaten_combo = ttk.Combobox(loader_frame, width=40, state="readonly")
        self.stammdaten_combo.pack(side=tk.LEFT, padx=5)
        
        # Populate combo
        stammdaten_list, _ = get_all_stammdaten_dict()
        self.stammdaten_combo['values'] = stammdaten_list
        
        ttk.Button(loader_frame, text="In Editor übernehmen", command=self.load_description_from_combo).pack(side=tk.LEFT, padx=5)

        # 3.2 Description
        desc_frame = ttk.Frame(self.editor_window)
        desc_frame.pack(fill='both', expand=True, pady=5, padx=10)
        ttk.Label(desc_frame, text="Beschreibung:").pack(anchor='w', padx=(0,5))
        self.description_text = tk.Text(desc_frame, height=5, width=60, font=("Segoe UI", 10))
        self.description_text.pack(side=tk.LEFT, fill='both', expand=True)
        desc_scroll = ttk.Scrollbar(desc_frame, command=self.description_text.yview)
        desc_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.description_text.config(yscrollcommand=desc_scroll.set)

        # 3.3 Amount & Buttons
        manual_frame = ttk.Frame(self.editor_window)
        manual_frame.pack(fill='x', pady=10, padx=10)
        ttk.Label(manual_frame, text="Manuelle Betragseingabe (€):").pack(side=tk.LEFT, padx=(0,5))
        self.amount_entry = ttk.Entry(manual_frame, width=10)
        self.amount_entry.pack(side=tk.LEFT, padx=(0,15))
        self.amount_entry.insert(0, "0.00")
        
        self.add_leistung_button = ttk.Button(manual_frame, text="Leistung Speichern", command=self.add_leistung_gui)
        self.add_leistung_button.pack(side=tk.LEFT, padx=10)
        ttk.Button(manual_frame, text="Editor Leeren", command=self.clear_editor).pack(side=tk.LEFT, padx=10)

    def _close_manual_editor_window(self):
        if hasattr(self, 'editor_window') and self.editor_window:
            if self.editor_window.winfo_exists():
                self.editor_window.destroy()
            self.editor_window = None
        
        # Referenzen auf Hauptfenster-Widgets wiederherstellen
        if hasattr(self, 'main_description_text'): self.description_text = self.main_description_text
        if hasattr(self, 'main_amount_entry'): self.amount_entry = self.main_amount_entry
        if hasattr(self, 'main_add_leistung_button'): self.add_leistung_button = self.main_add_leistung_button

    def clear_editor(self):
        self.description_text.delete("1.0", tk.END)
        self.amount_entry.delete(0, tk.END)
        self.amount_entry.insert(0, "0.00")
        if hasattr(self, 'description_text') and self.description_text.winfo_exists():
            self.description_text.delete("1.0", tk.END)
        if hasattr(self, 'amount_entry') and self.amount_entry.winfo_exists():
            self.amount_entry.delete(0, tk.END)
            self.amount_entry.insert(0, "0.00")
            
        self._reset_leistung_selection()
        # Reset button state if it was in edit mode
        self.add_leistung_button.config(text="Leistung Speichern", command=self.add_leistung_gui)
        if hasattr(self, 'add_leistung_button') and self.add_leistung_button.winfo_exists():
            self.add_leistung_button.config(text="Leistung Speichern", command=self.add_leistung_gui)
        self.selected_leistung_id = None

    def load_description_from_combo(self):
        selection = self.stammdaten_combo.get()
        if not selection: return
        
        if selection in self.stammdaten_betraege:
            betrag = self.stammdaten_betraege[selection]
            parts = selection.split(' - ', 1)
            desc = parts[1] if len(parts) > 1 else selection
            
            self.description_text.delete("1.0", tk.END)
            self.description_text.insert("1.0", desc)
            self.amount_entry.delete(0, tk.END)
            self.amount_entry.insert(0, f"{betrag:.2f}")

    def get_current_kilometergeld(self):
# ... (Rest der Klasse bleibt unverändert)
        """Ruft das Kilometergeld des aktuell ausgewählten Patienten ab."""
        if not self.patient_data:
            return 0.0
        # Kilometergeld ist Spalte 12, Index 11 in patient_data
        km_geld = self.patient_data[11] if len(self.patient_data) > 11 and self.patient_data[11] is not None else 0.0
        return km_geld

    def open_teamup_search(self):
        """
        Zusammengeführte Version:
        - Validierung (Patient & Leistungen)
        - Ladebalken während API-Abruf
        - Datumsbereichs-Automatik
        - ENTER-Support zum Ersetzen/Speichern
        """
        # --- 1. ORIGINAL VALIDIERUNG ---
        if not self.patient_data:
            messagebox.showwarning("Achtung", "Bitte wählen Sie zuerst einen Patienten im ersten Tab aus.")
            return
        if not self.selected_leistungs_kurznamen:
            messagebox.showwarning("Achtung", "Bitte wählen Sie zuerst im Hauptfenster mindestens eine Leistung (Button) aus, die den Terminen zugewiesen werden soll.")
            return

        # --- 2. DATUMSBERECHNUNG FÜR FILTER ---
        mode = CONFIG.get('AUTO_DATE_SELECTOR', 'Auto')
        first_day, last_day = None, None

        if mode == 'Manual':
            first_day = CONFIG.get('MANUAL_DATE_START', '')
            last_day = CONFIG.get('MANUAL_DATE_END', '')
            if not first_day or not last_day:
                first_day, last_day = None, None
        else:
            try:
                month = int(self.selected_invoice_month.get())
                year = int(self.selected_invoice_year.get())
                first_day = f"{year}-{month:02d}-01"
                last_day_num = calendar.monthrange(year, month)[1]
                last_day = f"{year}-{month:02d}-{last_day_num}"
            except Exception:
                first_day, last_day = None, None

        # NEU: Berechne UI-Defaults für die Anzeige (falls None, Standard +/- 30 Tage)
        ui_start_date = first_day
        ui_end_date = last_day
        if ui_start_date is None:
            ui_start_date = (datetime.date.today() - datetime.timedelta(days=30)).strftime('%Y-%m-%d')
        if ui_end_date is None:
            ui_end_date = (datetime.date.today() + datetime.timedelta(days=30)).strftime('%Y-%m-%d')

        initial_search_term = self.patient_data[2] if self.patient_data else ""

        # --- 3. LADEFENSTER (PROGRESSBAR) ---
        progress_win = tk.Toplevel(self.root)
        progress_win.title("API Abruf...")
        progress_win.geometry("300x120")
        progress_win.transient(self.root)
        progress_win.grab_set()
        
        # Zentrieren
        px = self.root.winfo_x() + (self.root.winfo_width() // 2) - 150
        py = self.root.winfo_y() + (self.root.winfo_height() // 2) - 60
        progress_win.geometry(f"+{px}+{py}")

        ttk.Label(progress_win, text=f"Lade Termine für:\n{initial_search_term}", justify="center").pack(pady=10)
        pb = ttk.Progressbar(progress_win, mode='indeterminate', length=200)
        pb.pack(pady=5)
        pb.start(15)
        progress_win.update()

        # Initialer API Abruf
        try:
            results = search_teamup_events(initial_search_term, start_date=first_day, end_date=last_day)
        finally:
            progress_win.destroy()

        # --- 4. DER SUCH-DIALOG (RE-DESIGNED) ---
        search_window = tk.Toplevel(self.root)
        search_window.title("Teamup Termin-Suche")
        search_window.geometry("650x600")
        search_window.grab_set()

        search_window.transient(self.root) 
        search_window.grab_set()

        ttk.Label(search_window, text="Suchbegriff (Name/Titel):").pack(pady=5, padx=10, anchor='w')
        search_entry = ttk.Entry(search_window, width=60)
        search_entry.pack(pady=5, padx=10)
        search_entry.insert(0, initial_search_term)
        
        # NEU: Datumsbereich Auswahl
        date_frame = ttk.Frame(search_window)
        date_frame.pack(pady=5, padx=10, anchor='w')
        
        ttk.Label(date_frame, text="Zeitraum von:").pack(side=tk.LEFT)
        start_date_entry = ttk.Entry(date_frame, width=12)
        start_date_entry.pack(side=tk.LEFT, padx=5)
        start_date_entry.insert(0, ui_start_date)
        
        ttk.Label(date_frame, text="bis:").pack(side=tk.LEFT)
        end_date_entry = ttk.Entry(date_frame, width=12)
        end_date_entry.pack(side=tk.LEFT, padx=5)
        end_date_entry.insert(0, ui_end_date)
        
        ttk.Label(date_frame, text="(YYYY-MM-DD)").pack(side=tk.LEFT, padx=5)
        
        # Treeview
        results_tree = ttk.Treeview(search_window, columns=('Titel', 'Datum', 'Von', 'Bis'), selectmode='extended', show='headings')
        results_tree.heading('Titel', text='Titel')
        results_tree.heading('Datum', text='Datum')
        results_tree.heading('Von', text='Von')
        results_tree.heading('Bis', text='Bis')
        results_tree.column('Datum', width=90, anchor='center')
        results_tree.column('Von', width=60, anchor='center')
        results_tree.column('Bis', width=60, anchor='center')
        results_tree.pack(pady=10, padx=10, expand=True, fill='both')

        # Daten einfüllen & alles markieren
        if results:
            for r in results:
                item = results_tree.insert('', tk.END, values=r)
                results_tree.selection_add(item)
        else:
            results_tree.insert('', tk.END, values=("keine Treffer", "", "", ""))

        # --- INTERNE FUNKTIONEN (ORIGINAL LOGIK) ---
        def perform_search(term=None, mode='standard'):
            search_t = term if term is not None else search_entry.get().strip()
            
            # NEU: Datum aus GUI lesen
            s_date = start_date_entry.get().strip()
            e_date = end_date_entry.get().strip()

            # Ladebalken anzeigen
            progress_win = tk.Toplevel(search_window)
            progress_win.title("Suche läuft...")
            progress_win.geometry("300x100")
            progress_win.transient(search_window)
            progress_win.grab_set()
            
            # Zentrieren
            try:
                px = search_window.winfo_x() + (search_window.winfo_width() // 2) - 150
                py = search_window.winfo_y() + (search_window.winfo_height() // 2) - 50
                progress_win.geometry(f"+{px}+{py}")
            except:
                pass

            ttk.Label(progress_win, text=f"Suche nach: {search_t}...", justify="center").pack(pady=10)
            pb = ttk.Progressbar(progress_win, mode='indeterminate', length=200)
            pb.pack(pady=5)
            pb.start(15)
            progress_win.update()

            try:
                # Nutzt auch hier den Datumsfilter
                res = search_teamup_events(search_t, start_date=s_date, end_date=e_date, mode=mode)
            finally:
                progress_win.destroy()

            results_tree.delete(*results_tree.get_children())
            if res:
                for r in res:
                    new_item = results_tree.insert('', tk.END, values=r)
                    results_tree.selection_add(new_item)
            else:
                results_tree.insert('', tk.END, values=("keine Treffer", "", "", ""))

        def _get_selected_events_and_validate():
            """Original Validierungs-Logik."""
            selected_items = results_tree.selection()
            if not selected_items:
                messagebox.showwarning("Achtung", "Bitte wählen Sie Termine aus der Liste aus.")
                return None

            valid_events = []
            for item in selected_items:
                vals = results_tree.item(item, 'values')
                if not vals or len(vals) < 4: continue
                title, date_str, time_from, time_to = vals
                # Zeit-Validierung
                if not (time_from and time_to and ':' in str(time_from)):
                    messagebox.showwarning("Fehler", f"Eintrag '{title}' enthält ungültige Zeitangaben.")
                    return None
                valid_events.append((title, date_str, time_from, time_to))
            return valid_events

        def add_selected_events():
            """Nutzt Original add_multiple_leistungen_from_teamup."""
            events = _get_selected_events_and_validate()
            if events:
                self.add_multiple_leistungen_from_teamup(events)
                search_window.destroy()

        def replace_selected_events():
            """Nutzt Original replace_all_leistungen_from_teamup."""
            events = _get_selected_events_and_validate()
            if events:
                self.replace_all_leistungen_from_teamup(events)
                search_window.destroy()

        def search_gemeinde():
            if not self.patient_data: return
            lastname = self.patient_data[2]
            
            # Setze Datum auf ganzes Jahr
            now = datetime.datetime.now()
            s_date = f"{now.year}-01-01"
            e_date = f"{now.year}-12-31"
            
            start_date_entry.delete(0, tk.END)
            start_date_entry.insert(0, s_date)
            end_date_entry.delete(0, tk.END)
            end_date_entry.insert(0, e_date)
            
            search_entry.delete(0, tk.END)
            search_entry.insert(0, lastname)
            
            perform_search(term=lastname, mode='gemeinde')

        # --- BUTTONS ---
        btn_frame = ttk.Frame(search_window)
        btn_frame.pack(pady=10)
        
        ttk.Button(btn_frame, text="Manuelle Suche", command=lambda: perform_search()).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Gemeinderechnung Suche", command=search_gemeinde).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="+ Hinzufügen", command=add_selected_events).pack(side=tk.LEFT, padx=5)
        save_btn = ttk.Button(btn_frame, text="Auswahl Speichern & Ersetzen (ENTER)", command=replace_selected_events)
        save_btn.pack(side=tk.LEFT, padx=5)
        search_window.after(100, lambda: self.focus_and_highlight(save_btn))

        # --- BINDINGS ---
        # ENTER im Suchfeld -> Neue Suche
        search_entry.bind('<Return>', lambda e: perform_search())
        
        # ENTER im Fenster (außerhalb Suchfeld) -> Ersetzen & Schließen
        search_window.bind('<Return>', lambda e: replace_selected_events() if search_window.focus_get() != search_entry else perform_search())
        
        # Doppelklick -> Hinzufügen (wie im Original)
        results_tree.bind('<Double-1>', lambda e: add_selected_events())
        # ESC -> Schließen
        search_window.bind('<Escape>', lambda e: search_window.destroy())

        # Fenster zentrieren
        search_window.update_idletasks()
        x = (search_window.winfo_screenwidth() // 2) - (search_window.winfo_width() // 2)
        y = (search_window.winfo_screenheight() // 2) - (search_window.winfo_height() // 2)
        search_window.geometry(f"+{x}+{y}")

    def load_leistung_stammdaten_buttons(self):
# ... (Rest der Klasse bleibt unverändert)
        """Lädt die Stammdaten und befüllt den Button-Bereich."""
        stammdaten_list, stammdaten_dict = get_all_stammdaten_dict()
        self.stammdaten_betraege = stammdaten_dict 
        
        # Update Combobox im Editor
        if hasattr(self, 'stammdaten_combo') and self.stammdaten_combo.winfo_exists():
            self.stammdaten_combo['values'] = stammdaten_list

        # FIX: Theme auf 'clam' setzen, damit Hintergrundfarben unterstützt werden
        # Windows-Standard-Themes (vista/xpnative) ignorieren oft background-Farben
        if self.ttk_style.theme_use() != 'clam':
            self.ttk_style.theme_use('clam')

        # Style definieren (Grüner Hintergrund, fetter Text)
        self.ttk_style.configure('Selected.TButton', 
                            background='#90EE90', # Hellgrün
                            foreground='black', 
                            bordercolor='#006400', # Dunkelgrüner Rand (Highlight)
                            lightcolor='#90EE90',
                            darkcolor='#90EE90',
                            borderwidth=2,
                            font=('Helvetica', 9, 'bold'))
        
        self.ttk_style.map('Selected.TButton',
            background=[('active', '#7ccd7c'), ('pressed', '#66bb66')],
            bordercolor=[('active', '#004d00'), ('pressed', '#004d00')],
        )


        # Lösche vorhandene Buttons
        for widget in self.leistung_button_frame.winfo_children():
            widget.destroy()

        # Erstelle neue Buttons
        for i, item in enumerate(stammdaten_list):
            kurzname = item.split(' - ')[0]
            betrag = stammdaten_dict[item]
            
            btn = ttk.Button(self.leistung_button_frame, 
                             text=f"{kurzname} (€{betrag:.2f})", 
                             command=lambda k=kurzname: self.toggle_leistung_selection(k))
            
            btn.kurzname = kurzname
            btn.is_selected = False
            btn.grid(row=i // 5, column=i % 5, padx=5, pady=5, sticky='w')
            
        self.leistung_button_frame.update_idletasks()
        self.leistung_canvas.config(scrollregion=self.leistung_canvas.bbox("all"))

    def load_and_select_last_leistungen(self):
        # ... (Rest der Klasse bleibt unverändert)
        """Lädt die zuletzt gespeicherte Leistungsauswahl für den Patienten und wählt die Buttons."""
        if not self.patient_data:
            return 
            
        # 1. Zuerst bestehende Auswahl zurücksetzen
        self._reset_leistung_selection()

        patient_id = self.patient_data[0]
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        # 2. Letzte Auswahl abrufen (Index 12)
        cursor.execute("SELECT last_selected_kurznamen FROM patienten WHERE id = ?", (patient_id,))
        row = cursor.fetchone()
        kurznamen_str = row[0] if row else ""
        conn.close()
        
        if not kurznamen_str:
            return

        kurznamen_set = set(kurznamen_str.split(','))
        
        # 3. Buttons entsprechend auswählen
        for kurzname_to_select in kurznamen_set:
            # Finde den Button und wähle ihn aus
            for widget in self.leistung_button_frame.winfo_children():
                if hasattr(widget, 'kurzname') and widget.kurzname == kurzname_to_select:
                    self.selected_leistungs_kurznamen.add(kurzname_to_select)
                    widget.is_selected = True
                    widget.config(style='Selected.TButton')
                    break

    def toggle_leistung_selection(self, kurzname):
        button_ref = None
        for widget in self.leistung_button_frame.winfo_children():
            if hasattr(widget, 'kurzname') and widget.kurzname == kurzname:
                button_ref = widget
                break
        
        if not button_ref: return

        if kurzname in self.selected_leistungs_kurznamen:
            self.selected_leistungs_kurznamen.remove(kurzname)
            button_ref.config(style='TButton') # Zurück zum Standard
        else:
            self.selected_leistungs_kurznamen.add(kurzname)
            button_ref.config(style='Selected.TButton') # Grün markieren

    def add_leistung_gui(self):
        # ... (Rest der Klasse bleibt unverändert)
        """Fügt neue Leistung(en) in die DB ein, nun mit Uhrzeit und Kilometergeld-Zuschlag."""
        if not self.patient_data:
            messagebox.showwarning("Warnung", "Bitte wählen Sie zuerst einen Patienten aus.")
            return

        patient_id = self.patient_data[0]
        datum_str = self.date_entry.get().strip()
        time_from_str = self.time_from_entry.get().strip()
        time_to_str = self.time_to_entry.get().strip()
        
        manual_betrag, use_manual_override, km_geld = self._get_leistung_insertion_params()
        custom_desc = self.description_text.get("1.0", tk.END).strip()

        try:
            datetime.datetime.strptime(datum_str, '%d.%m.%Y')
            if len(time_from_str) < 5 or len(time_to_str) < 5 or ":" not in time_from_str:
                raise ValueError("Uhrzeit muss im Format HH:MM angegeben werden.")
        except ValueError as e:
            messagebox.showerror("Fehler", f"Ungültiges Datums- oder Zeitformat: {e}")
            return
        
        success_count = 0
        
        if use_manual_override:
            # Nur eine Leistung (manuelle) hinzufügen
            kurzname = "Manuell"
            standard_betrag = manual_betrag
            if self.add_leistung_to_db(patient_id, datum_str, time_from_str, time_to_str, kurzname, standard_betrag, manual_betrag, use_manual_override, km_geld, custom_description=custom_desc):
                success_count = 1
        
        else:
            # Leistungen basierend auf ausgewählten Buttons hinzufügen
            if not self.selected_leistungs_kurznamen:
                messagebox.showwarning("Achtung", "Bitte wählen Sie mindestens eine Leistung aus der Liste aus oder geben Sie einen manuellen Betrag ein.")
                return

            for kurzname in self.selected_leistungs_kurznamen:
                # Finde den vollen Stammdaten-Key für den Betrag
                stammdaten_key = [k for k in self.stammdaten_betraege.keys() if k.startswith(kurzname + ' -')]
                if not stammdaten_key:
                    logging.warning(f"Stammdaten für '{kurzname}' nicht gefunden. Überspringe.")
                    continue
                
                standard_betrag = self.stammdaten_betraege[stammdaten_key[0]]
                
                if self.add_leistung_to_db(patient_id, datum_str, time_from_str, time_to_str, kurzname, standard_betrag, manual_betrag, use_manual_override, km_geld):
                    success_count += 1
            
            # Speichere die aktuelle Auswahl für diesen Patienten
            save_last_selected_leistungen(patient_id, self.selected_leistungs_kurznamen)


        if success_count > 0:
            self.set_status(f"✅ {success_count} Leistung(en) erfolgreich hinzugefügt.")
            self.description_text.delete("1.0", tk.END) # Textbox leeren
            if hasattr(self, 'description_text') and self.description_text.winfo_exists():
                self.description_text.delete("1.0", tk.END) # Textbox leeren
            self.update_leistung_list()
            
            # Editor schließen, wenn er offen ist
            self._close_manual_editor_window()
        elif success_count == 0 and not use_manual_override:
            messagebox.showwarning("Achtung", "Es konnten keine neuen Leistungen hinzugefügt werden (Prüfen Sie, ob Stammdaten fehlen).")

    def add_multiple_leistungen_from_teamup(self, events_list):
# ... (Rest der Klasse bleibt unverändert)
        """Fügt mehrere Leistungen (Teamup) hinzu und speichert die Auswahl."""
        if not self.patient_data or not self.selected_leistungs_kurznamen:
            messagebox.showwarning("Fehler", "Kein Patient oder keine Leistung ausgewählt. Vorgang abgebrochen.")
            return

        patient_id = self.patient_data[0]
        # Hier wird die Kernlogik aufgerufen
        total_success_count = self._insert_multiple_leistungen(patient_id, events_list) 

        if total_success_count > 0:
            # Speichere die aktuelle Auswahl für diesen Patienten
            save_last_selected_leistungen(patient_id, self.selected_leistungs_kurznamen)
            self.set_status(f"{total_success_count} Leistung(en) für Patient {self.patient_data[2]} erfolgreich hinzugefügt.")
            self.update_leistung_list()
            self.root.focus_force()
        # Keine MessageBox bei 0, da das System das intern loggen kann.

    def replace_all_leistungen_from_teamup(self, events_list):
# ... (Rest der Klasse bleibt unverändert)
        """Löscht alle bestehenden Leistungen des Patienten und fügt die ausgewählten Teamup-Termine als neue Leistungen ein."""
        if not self.patient_data or not self.selected_leistungs_kurznamen:
            logging.warning("Kein Patient oder keine Leistung ausgewählt. Vorgang abgebrochen.")
            return

        patient_id = self.patient_data[0]
        patient_name = f"{self.patient_data[1]} {self.patient_data[2]}"

        # 1. Lösche alle bestehenden Leistungen
        if not self._delete_all_patient_leistungen(patient_id):
            messagebox.showerror("Fehler", f"FEHLER: Fehler beim Löschen bestehender Leistungen für Patient {patient_name}. Neue Termine wurden NICHT eingefügt.")
            return

        # 2. Füge neue Leistungen ein
        insertion_success_count = self._insert_multiple_leistungen(patient_id, events_list)

        if insertion_success_count > 0:
            # Speichere die aktuelle Auswahl für diesen Patienten
            save_last_selected_leistungen(patient_id, self.selected_leistungs_kurznamen)
            self.set_status(f"{insertion_success_count} Leistung(en) für Patient {patient_name} erfolgreich ERSETZT.")
            self.update_leistung_list()
            self.root.focus_force()
        else:
            messagebox.showwarning("Achtung", "Es konnten keine neuen Leistungen hinzugefügt werden (nach dem Löschen).")
            logging.info("Es konnten keine neuen Leistungen hinzugefügt werden (nach dem Löschen).")

    def update_leistung_list(self):
        # ... (Rest der Klasse bleibt unverändert)
        """Aktualisiert die Liste der Leistungen im Treeview."""
        if not self.patient_data:
            self.leistung_tree.delete(*self.leistung_tree.get_children())
            self.summary_label.config(text="Gesamtsumme: €0.00 | Nicht abgerechnete Leistungen: 0")
            return

        patient_id = self.patient_data[0]
        leistungen = get_patient_leistungen(patient_id)

        self.leistung_tree.delete(*self.leistung_tree.get_children())
        total_sum = 0.0

        for leistung in leistungen:
            leistung_id, datum_db, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag = leistung
            datum_formatiert = datetime.datetime.strptime(datum_db, '%Y-%m-%d').strftime('%d.%m.%Y')
            
            # Darstellung des Betrags (bereits inkl. KM-Geld)
            betrag_str = f"€{einzelbetrag:.2f}"
            total_sum += einzelbetrag

            self.leistung_tree.insert('', tk.END, iid=leistung_id, values=(
                leistung_id, datum_formatiert, uhrzeit_von, uhrzeit_bis, beschreibung, betrag_str
            ))

        self.summary_label.config(text=f"Gesamtsumme: €{total_sum:.2f} | Nicht abgerechnete Leistungen: {len(leistungen)}")

    def select_leistung_for_edit(self, event):
        # ... (Rest der Klasse bleibt unverändert)
        """Speichert die ID der ausgewählten Leistung."""
        selected_item = self.leistung_tree.focus()
        if selected_item:
            self.selected_leistung_id = self.leistung_tree.item(selected_item)['values'][0]
        else:
            self.selected_leistung_id = None
        
        # Setze den Button zurück, falls nichts ausgewählt ist
        if not self.selected_leistung_id:
            if hasattr(self, 'add_leistung_button') and self.add_leistung_button.winfo_exists():
                self.add_leistung_button.config(text="Leistung Hinzufügen (Manuell/Auswahl)", command=self.add_leistung_gui)


    def load_leistung_for_edit(self):
        # ... (Rest der Klasse bleibt unverändert)
        """Lädt die ausgewählte Leistung in die Eingabefelder zum Bearbeiten."""
        if not hasattr(self, 'selected_leistung_id') or not self.selected_leistung_id:
            messagebox.showwarning("Achtung", "Bitte wählen Sie eine Leistung aus der Liste aus.")
            return

        leistung_id = self.selected_leistung_id
        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        cursor.execute("""
        SELECT datum, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag 
        FROM leistungen 
        WHERE id = ? 
        """, (leistung_id,))
        res = cursor.fetchone()
        conn.close()

        if not res:
            messagebox.showerror("Fehler", f"Leistung ID {leistung_id} konnte nicht in der Datenbank gefunden werden.")
            return

        # [0]datum_db, [1]uhrzeit_von, [2]uhrzeit_bis, [3]beschreibung, [4]einzelbetrag (inkl. KM-Geld)
        # Um den Basis-Betrag anzuzeigen, müssen wir das KM-Geld wieder abziehen:
        km_geld = self.get_current_kilometergeld()
        basis_betrag = res[4] - km_geld

        datum_formatiert = datetime.datetime.strptime(res[0], '%Y-%m-%d').strftime('%d.%m.%Y')
        uhrzeit_von = res[1]
        uhrzeit_bis = res[2]
        beschreibung = res[3] 

        betrag_str = f"{basis_betrag:.2f}"

        # Öffne den Editor automatisch ZUERST, damit die Referenzen (amount_entry, description_text) auf das Fenster zeigen
        self.open_manual_editor()

        self.date_entry.delete(0, tk.END)
        self.date_entry.insert(0, datum_formatiert)
        self.time_from_entry.delete(0, tk.END)
        self.time_from_entry.insert(0, uhrzeit_von)
        self.time_to_entry.delete(0, tk.END)
        self.time_to_entry.insert(0, uhrzeit_bis)
        self.amount_entry.delete(0, tk.END)
        self.amount_entry.insert(0, betrag_str)
        self.description_text.delete("1.0", tk.END)
        self.description_text.insert("1.0", beschreibung)
        
        
        # NEU: Beim Bearbeiten alle Buttons abwählen, da der Betrag manuell gesetzt wird
        # self._reset_leistung_selection() 
        self.use_km_money_var.set(True) # Reset auf Standard (da KM-Geld aus Anzeige herausgerechnet wurde)

        # Button-Funktion auf Update umstellen
        self.add_leistung_button.config(text=f"Leistung Aktualisieren (ID: {leistung_id})", command=lambda: self.update_leistung_gui(leistung_id))
        logging.info(f"Leistung ID {leistung_id} zum Bearbeiten geladen. Basisbetrag (€{basis_betrag:.2f}) angezeigt.")


    def update_leistung_gui(self, leistung_id):
        # ... (Rest der Klasse bleibt unverändert)
        """Aktualisiert eine bestehende Leistung, nun mit Uhrzeit und Kilometergeld-Zuschlag."""
        datum_str = self.date_entry.get().strip()
        time_from_str = self.time_from_entry.get().strip()
        time_to_str = self.time_to_entry.get().strip()
        betrag_str = self.amount_entry.get().strip().replace(',', '.') 
        beschreibung = self.description_text.get("1.0", tk.END).strip()

        if self.use_km_money_var.get():
            km_geld = self.get_current_kilometergeld()
        else:
            km_geld = 0.0
        
        try:
            datum_db = datetime.datetime.strptime(datum_str, '%d.%m.%Y').strftime('%Y-%m-%d')
            betrag_basis = float(betrag_str)
            end_betrag = betrag_basis + km_geld

            if len(time_from_str) < 5 or len(time_to_str) < 5 or ":" not in time_from_str:
                raise ValueError("Uhrzeit muss im Format HH:MM angegeben werden.")

            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            
            # Aktualisiere die Beschreibung, da es sich nun um eine manuelle Bearbeitung handelt
            if not beschreibung:
                beschreibung = f"Manuelle Korrektur (ID {leistung_id})"

            cursor.execute("""
            UPDATE leistungen
            SET datum=?, uhrzeit_von=?, uhrzeit_bis=?, beschreibung=?, einzelbetrag=?
            WHERE id=?
            """, (datum_db, time_from_str, time_to_str, beschreibung, end_betrag, leistung_id))
            conn.commit()

            self.set_status(f"Leistung ID {leistung_id} erfolgreich aktualisiert.")
            self.update_leistung_list()
            
            # Setze den Button zurück auf Hinzufügen-Modus
            self.add_leistung_button.config(text="Leistung Hinzufügen (Manuell/Auswahl)", command=self.add_leistung_gui)
            self.selected_leistung_id = None
            self.description_text.delete("1.0", tk.END)
            
            # Editor schließen, wenn er offen ist
            self._close_manual_editor_window()
            
        except ValueError as e:
            messagebox.showerror("Fehler", f"Ungültiges Datums-, Zeit- oder Betragsformat: {e}")
        except Exception as e:
            messagebox.showerror("Fehler", f"Datenbankfehler beim Aktualisieren: {e}")
            conn.close()

    def delete_leistung_gui(self):
# ... (Rest der Klasse bleibt unverändert)
        """Löscht die aktuell ausgewählte Leistung."""
        if not hasattr(self, 'selected_leistung_id') or not self.selected_leistung_id:
            messagebox.showwarning("Achtung", "Bitte wählen Sie die zu löschende Leistung aus der Liste aus.")
            return

        leistung_id = self.selected_leistung_id
        if messagebox.askyesno("Bestätigen", f"Sind Sie sicher, dass Sie die Leistung ID {leistung_id} löschen möchten?"):
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            try:
                cursor.execute("DELETE FROM leistungen WHERE id = ?", (leistung_id,))
                conn.commit()
                logging.info(f"Leistung ID {leistung_id} erfolgreich gelöscht.")
                self.update_leistung_list()
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Löschen: {e}")
            finally:
                conn.close()

    def delete_all_leistungen_gui(self):
# ... (Rest der Klasse bleibt unverändert)
        """Löscht ALLE Leistungen des aktuell ausgewählten Patienten."""
        if not self.patient_data:
            messagebox.showwarning("Achtung", "Kein Patient ausgewählt.")
            return

        patient_id = self.patient_data[0]
        patient_name = f"{self.patient_data[1]} {self.patient_data[2]}"
        if messagebox.askyesno("WARNUNG", f"Sind Sie sicher, dass Sie ALLE Leistungen für Patient '{patient_name}' (ID: {patient_id}) löschen möchten? Dieser Schritt kann nicht rückgängig gemacht werden."):
            if self._delete_all_patient_leistungen(patient_id):
                logging.info(f"Alle Leistungen für Patient '{patient_name}' wurden gelöscht.")
                self.update_leistung_list()

    # --- 4. Stammdaten Leistungen Tab ---
    def setup_stammdaten_tab(self, tab):
        # Clear existing widgets if any
        for widget in tab.winfo_children():
            widget.destroy()

        # --- Input Bereich (Oben) ---
        input_frame = ttk.LabelFrame(tab, text="Leistung bearbeiten / neu anlegen", padding=15)
        input_frame.pack(fill='x', padx=10, pady=10)

        # Zeile 1: Kurzname und Betrag
        row1 = ttk.Frame(input_frame)
        row1.pack(fill='x', pady=5)
        
        ttk.Label(row1, text="Kurzname (Eindeutig):").pack(side=tk.LEFT)
        self.stammdaten_kurzname_entry = ttk.Entry(row1, width=25)
        self.stammdaten_kurzname_entry.pack(side=tk.LEFT, padx=(5, 20))
        
        ttk.Label(row1, text="Standard Betrag (€):").pack(side=tk.LEFT)
        self.stammdaten_betrag_entry = ttk.Entry(row1, width=15)
        self.stammdaten_betrag_entry.pack(side=tk.LEFT, padx=5)

        # Schnellauswahl Buttons
        def _set_quick_amt(val):
            self.stammdaten_betrag_entry.delete(0, tk.END)
            self.stammdaten_betrag_entry.insert(0, f"{val:.2f}")

        # Werte aus Config laden
        quick_amounts_str = CONFIG.get('QUICK_AMOUNTS', "33, 35, 57, 83")
        try:
            # String "33, 57, 80" -> Liste [33.0, 57.0, 80.0]
            quick_amounts = [float(x.strip()) for x in quick_amounts_str.split(',') if x.strip()]
        except ValueError:
            quick_amounts = [33, 35, 57, 83] # Fallback bei Tippfehlern

        for val in quick_amounts:
            # Anzeige ohne Nachkommastellen, wenn es eine ganze Zahl ist (z.B. 57 statt 57.0)
            label_text = f"€{int(val)}" if val.is_integer() else f"€{val}"
            ttk.Button(row1, text=label_text, width=5, command=lambda v=val: _set_quick_amt(v)).pack(side=tk.LEFT, padx=2)

        # Zeile 2: Beschreibung Label
        ttk.Label(input_frame, text="Beschreibung (für Honorarnote):").pack(anchor='w', pady=(15, 5))

        # Zeile 3: Beschreibung Textbox (Groß)
        desc_frame = ttk.Frame(input_frame)
        desc_frame.pack(fill='x', pady=0)
        
        self.stammdaten_desc_text = tk.Text(desc_frame, height=6, font=("Segoe UI", 10))
        self.stammdaten_desc_text.pack(side=tk.LEFT, fill='x', expand=True)
        
        scrollbar = ttk.Scrollbar(desc_frame, command=self.stammdaten_desc_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill='y')
        self.stammdaten_desc_text.config(yscrollcommand=scrollbar.set)

        # Zeile 4: Buttons
        btn_frame = ttk.Frame(input_frame)
        btn_frame.pack(fill='x', pady=(15, 0))
        
        ttk.Button(btn_frame, text="💾 Speichern / Aktualisieren", command=self.save_stammdaten).pack(side=tk.RIGHT)
        ttk.Button(btn_frame, text="Felder leeren", command=self.clear_stammdaten_fields).pack(side=tk.RIGHT, padx=10)

        # --- Liste Bereich (Mitte) ---
        list_frame = ttk.LabelFrame(tab, text="Vorhandene Leistungen (Stammdaten)", padding=10)
        list_frame.pack(fill='both', expand=True, padx=10, pady=5)

        self.stammdaten_listbox = tk.Listbox(list_frame, height=10, font=("Segoe UI", 10))
        self.stammdaten_listbox.pack(side=tk.LEFT, fill='both', expand=True)
        
        list_scroll = ttk.Scrollbar(list_frame, command=self.stammdaten_listbox.yview)
        list_scroll.pack(side=tk.RIGHT, fill='y')
        self.stammdaten_listbox.config(yscrollcommand=list_scroll.set)
        
        self.stammdaten_listbox.bind('<<ListboxSelect>>', self.select_stammdaten_from_list)
        self.stammdaten_listbox.bind('<Button-3>', self._show_stammdaten_context_menu)

        # --- Controls (Unten) ---
        control_frame = ttk.Frame(tab)
        control_frame.pack(fill='x', padx=10, pady=10)
        ttk.Button(control_frame, text="🗑️ Ausgewählte Leistung Löschen", command=self.delete_stammdaten).pack(side=tk.LEFT)
        ttk.Button(control_frame, text="📁 Ausgewählte Leistung Archivieren", command=self.archive_stammdaten).pack(side=tk.LEFT, padx=10)

        self.update_stammdaten_list()

    def clear_stammdaten_fields(self):
        self.stammdaten_kurzname_entry.delete(0, tk.END)
        self.stammdaten_betrag_entry.delete(0, tk.END)
        self.stammdaten_desc_text.delete("1.0", tk.END)

    def update_stammdaten_list(self):
        # ... (Rest der Klasse bleibt unverändert)
        """Aktualisiert die Liste der Stammdaten."""
        self.stammdaten_listbox.delete(0, tk.END)
        stammdaten_list, stammdaten_dict = get_all_stammdaten_dict()
        for item in stammdaten_list:
            betrag = stammdaten_dict[item]
            self.stammdaten_listbox.insert(tk.END, f"{item} (Standard: €{betrag:.2f})")
            
        self.load_leistung_stammdaten_buttons() # Update auch die Buttons

    def _show_stammdaten_context_menu(self, event):
        """Zeigt Kontextmenü für Stammdaten-Liste (nur Kopieren)."""
        try:
            index = self.stammdaten_listbox.nearest(event.y)
            if index == -1: return
            
            # Selektiere die Zeile unter der Maus
            self.stammdaten_listbox.selection_clear(0, tk.END)
            self.stammdaten_listbox.selection_set(index)
            self.stammdaten_listbox.activate(index)
            
            menu = tk.Menu(self.root, tearoff=0)
            menu.add_command(label="Kopieren", command=lambda: self._copy_stammdaten_description(index))
            menu.tk_popup(event.x_root, event.y_root)
        except Exception:
            pass

    def _copy_stammdaten_description(self, index):
        try:
            item_text = self.stammdaten_listbox.get(index)
            # Format: "Kurzname - Beschreibung (Standard: €...)"
            full_desc = item_text.split(' (Standard: ')[0]
            parts = full_desc.split(' - ', 1)
            beschreibung = parts[1] if len(parts) > 1 else ""
            
            if beschreibung:
                self.root.clipboard_clear()
                self.root.clipboard_append(beschreibung)
                self.set_status(f"Beschreibung kopiert: {beschreibung}")
        except Exception:
            pass

    def select_stammdaten_from_list(self, event):
        # ... (Rest der Klasse bleibt unverändert)
        """Lädt die ausgewählte Stammdatenleistung in die Eingabefelder."""
        selection = self.stammdaten_listbox.curselection()
        if selection:
            item_text = self.stammdaten_listbox.get(selection[0])
            # Extrahiere Kurzname, Beschreibung und Betrag
            match = item_text.split(' (Standard: ')
            full_desc = match[0]
            betrag_str = match[1].replace('€', '').replace(')', '')
            
            parts = full_desc.split(' - ', 1)
            kurzname = parts[0]
            beschreibung = parts[1] if len(parts) > 1 else ""
            
            self.stammdaten_kurzname_entry.delete(0, tk.END)
            self.stammdaten_kurzname_entry.insert(0, kurzname)
            self.stammdaten_betrag_entry.delete(0, tk.END)
            self.stammdaten_betrag_entry.insert(0, betrag_str)
            self.stammdaten_desc_text.delete("1.0", tk.END)
            self.stammdaten_desc_text.insert("1.0", beschreibung)


    def save_stammdaten(self):
        # ... (Rest der Klasse bleibt unverändert)
        """Speichert oder aktualisiert Stammdaten."""
        kurzname = self.stammdaten_kurzname_entry.get().strip()
        betrag_str = self.stammdaten_betrag_entry.get().strip().replace(',', '.')
        beschreibung = self.stammdaten_desc_text.get("1.0", tk.END).strip()

        if not kurzname or not beschreibung or not betrag_str:
            messagebox.showwarning("Achtung", "Alle Felder müssen ausgefüllt sein.")
            return

        try:
            betrag = float(betrag_str)
        except ValueError:
            messagebox.showwarning("Achtung", "Betrag muss eine gültige Zahl sein.")
            return

        conn = sqlite3.connect(DATABASE_NAME)
        cursor = conn.cursor()
        
        try:
            # UPDATE (wenn kurzname schon existiert)
            cursor.execute("""
            UPDATE stammdaten_leistungen 
            SET beschreibung = ?, standard_betrag = ?
            WHERE kurzname = ?
            """, (beschreibung, betrag, kurzname))
            
            if cursor.rowcount == 0:
                # INSERT
                cursor.execute("""
                INSERT INTO stammdaten_leistungen (kurzname, beschreibung, standard_betrag)
                VALUES (?, ?, ?)
                """, (kurzname, beschreibung, betrag))
                messagebox.showinfo("Erfolg", f"Neue Leistung '{kurzname}' erfolgreich erstellt.")
            else:
                messagebox.showinfo("Erfolg", f"Leistung '{kurzname}' erfolgreich aktualisiert.")
                
            conn.commit()
            self.update_stammdaten_list()
            
            self.clear_stammdaten_fields()
                
        except sqlite3.IntegrityError:
            messagebox.showerror("Fehler", f"Kurzname '{kurzname}' existiert bereits. Bitte ändern Sie den Kurznamen.")
        except Exception as e:
            messagebox.showerror("Fehler", f"Datenbankfehler: {e}")
            
        conn.close()

    def archive_stammdaten(self):
        """Archiviert die ausgewählte Stammdatenleistung."""
        selection = self.stammdaten_listbox.curselection()
        if not selection:
            messagebox.showwarning("Achtung", "Bitte wählen Sie eine Leistung aus der Liste aus.")
            return

        item_text = self.stammdaten_listbox.get(selection[0])
        kurzname = item_text.split(' - ')[0]

        if messagebox.askyesno("Bestätigen", f"Sind Sie sicher, dass Sie die Stammdatenleistung '{kurzname}' archivieren möchten?\nSie wird aus der aktiven Liste entfernt."):
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            try:
                cursor.execute("UPDATE stammdaten_leistungen SET is_archived = 1 WHERE kurzname = ?", (kurzname,))
                conn.commit()
                messagebox.showinfo("Erfolg", f"Stammdatenleistung '{kurzname}' erfolgreich archiviert.")
                self.update_stammdaten_list()
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Archivieren: {e}")
            finally:
                conn.close()

    def delete_stammdaten(self):
# ... (Rest der Klasse bleibt unverändert)
        """Löscht die ausgewählte Stammdatenleistung."""
        selection = self.stammdaten_listbox.curselection()
        if not selection:
            messagebox.showwarning("Achtung", "Bitte wählen Sie eine Leistung aus der Liste aus.")
            return

        item_text = self.stammdaten_listbox.get(selection[0])
        kurzname = item_text.split(' - ')[0]

        if messagebox.askyesno("Bestätigen", f"Sind Sie sicher, dass Sie die Stammdatenleistung '{kurzname}' löschen möchten?"):
            conn = sqlite3.connect(DATABASE_NAME)
            cursor = conn.cursor()
            try:
                cursor.execute("DELETE FROM stammdaten_leistungen WHERE kurzname = ?", (kurzname,))
                conn.commit()
                messagebox.showinfo("Erfolg", f"Stammdatenleistung '{kurzname}' erfolgreich gelöscht.")
                self.update_stammdaten_list()
            except Exception as e:
                messagebox.showerror("Fehler", f"Fehler beim Löschen: {e}")
            finally:
                conn.close()

# --- START DER ANWENDUNG ---
if __name__ == "__main__":
    start_gui()