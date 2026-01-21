# leprendix/services/doc_generator.py
import os
import datetime
import logging
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_UNDERLINE

from leprendix.core.paths import TEMPLATE_PATH
from leprendix.core.config_loader import CONFIG
from leprendix.db.queries import get_patient_leistungen_for_template

def fill_template(patient_id, patient_data_tuple, template_data, add_gemeinde_block, ausstellungs_datum, ueberweisung=True): 
    """Füllt die Word-Vorlage mit den Patientendaten und Leistungen und speichert sie."""
    
    # --- Einrückungs-Konstanten ---
    LEISTUNG_INDENT_SIZE = Pt(70)  
    SPACE_AFTER_LEISTUNG_BLOCK = Pt(12) 

    # Unpacking für 13 Elemente (ohne ueberweisung in DB)
    _, vorname, nachname, strasse, hausnummer, adresszusatz, plz, ort, anrede, versicherungsnummer, diagnose, kilometergeld, _ = patient_data_tuple
    
    leistungen_liste = get_patient_leistungen_for_template(patient_id)
    
    heute = ausstellungs_datum # Verwendet das übergebene Datum
    
    # Sicherheitscheck: Falls ueberweisung None ist, Standard auf True setzen
    if ueberweisung is None:
        ueberweisung = True
    logging.info(f"Generiere Rechnung. Überweisung-Modus: {ueberweisung}")

    try:
        document = Document(TEMPLATE_PATH)
    except FileNotFoundError:
        raise FileNotFoundError(f"Die Vorlagendatei '{TEMPLATE_PATH}' wurde nicht gefunden.")
    
    total_betrag = 0.0
    
    if not document.paragraphs:
        raise ValueError("Word-Vorlage enthält keine Paragraphen.")

    invoice_number_paragraph = document.paragraphs[0]
            
    # Statische Platzhalter
    replacements = {
        '{{Rechnungsnummer}}': template_data['BHAG_NUMMER'], 
        '{{Anrede}}': anrede, 
        '{{Nachname}}': nachname,
        '{{Vorname}}': vorname,
        '{{Straße}}': strasse,
        '{{Hausnummer}}': hausnummer,
        '{{Adresszusatz}}': adresszusatz or '',
        '{{Postleitzahl}}': plz,
        '{{Stadt}}': ort,
        '{{Versicherungsnummer}}': versicherungsnummer,
        '{{Datum_Austellung}}': heute,
        '{{Diagnose}}': diagnose
    }
    
    # --- 0. Globale Formatierung (Vorbereitung und HONORARNOTE Sonderfall) ---
    
    HONORARNOTE_PARAGRAPH = None
    ORT_DATUM_PARAGRAPH = None
    DIAGNOSE_PLACEHOLDER_PARAGRAPH = None 
    
    for i, p in enumerate(document.paragraphs):
        
        if 'H O N O R A R N O T E' in p.text:
            HONORARNOTE_PARAGRAPH = p
            
        if 'Mödling, {{Datum_Austellung}}' in p.text:
            ORT_DATUM_PARAGRAPH = p
            
        if '{{Diagnose}}' in p.text:
            DIAGNOSE_PLACEHOLDER_PARAGRAPH = p
            
        # Setze die Standardgröße 12pt global
        for run in p.runs:
            run.font.size = Pt(12) 
            
    # Spezielle Formatierung für HONORARNOTE anwenden
    if HONORARNOTE_PARAGRAPH:
        for run in HONORARNOTE_PARAGRAPH.runs:
            run.font.size = Pt(18)
        HONORARNOTE_PARAGRAPH.paragraph_format.space_after = Pt(10) 
        HONORARNOTE_PARAGRAPH.paragraph_format.space_before = Pt(0) 
        HONORARNOTE_PARAGRAPH.paragraph_format.line_spacing = 1.0 


    # --- 1. Rechnungsnummer ersetzen und Gemeinde-Block einfügen ---
    
    if '{{Rechnungsnummer}}' in invoice_number_paragraph.text:
        invoice_number_paragraph.text = invoice_number_paragraph.text.replace('{{Rechnungsnummer}}', template_data['BHAG_NUMMER'])
        del replacements['{{Rechnungsnummer}}']
        
    if add_gemeinde_block:
        # 1. Ziel-Index festlegen (direkt unter der Rechnungsnummer)
        target_index = 1 
        try:
            target_p = document.paragraphs[target_index]
        except IndexError:
            target_p = document.paragraphs[-1]

        # 2. Einen neuen Absatz genau dort einfügen
        new_p = target_p.insert_paragraph_before()
        new_p.paragraph_format.line_spacing = 1.0
        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.space_after = Pt(0)

        # 3. Den Inhalt zusammenbauen:
        # Eine Leerzeile oben (\n)
        # Dann die Adresse
        # Dann vier Leerzeilen unten (\n\n\n\n)
        full_text = (
            "\n\n\n" +
            "Gemeinde Wiener Neudorf\n" +
            "Europaplatz 2\n" +
            "2351 Wiener Neudorf" +
            "\n\n\n\n\n"
        )

        # 4. Den Text dem Absatz hinzufügen und fett/12pt machen
        run = new_p.add_run(full_text)
        run.bold = True
        run.font.size = Pt(12)

    # --- 2. Dynamische Leistungsblock-Logik und restliche Ersetzung (mit Seitenumbruch-Schutz) ---
    
    start_tag = '{{LEISTUNGSBLOCK_START}}'
    end_tag = '{{LEISTUNGSBLOCK_ENDE}}'
    
    block_start_paragraph = None
    block_end_paragraph = None
    block_paragraphs = []
    
    in_block = False 
    
    # Text-Ersetzung und Block-Suche in einem Durchlauf
    for p in document.paragraphs:
        
        if p is HONORARNOTE_PARAGRAPH or p is invoice_number_paragraph or p is DIAGNOSE_PLACEHOLDER_PARAGRAPH:
            continue
            
        # Statische Platzhalter ersetzen
        for key, value in replacements.items():
            if key in p.text:
                p.text = p.text.replace(key, value)
        
        # Block-Logik: Musterblock finden
        if start_tag in p.text:
            block_start_paragraph = p
            in_block = True
            continue
        
        if end_tag in p.text:
            block_end_paragraph = p
            break
            
        if in_block:
            block_paragraphs.append(p)

    # Generieren des gesamten Leistungs-Textes
    gesamt_leistungs_text = ""
    
    if block_start_paragraph and block_end_paragraph:
        
        template_text = '\n'.join([p.text for p in block_paragraphs])
        
        for datum_db, uhrzeit_von, uhrzeit_bis, beschreibung, einzelbetrag in leistungen_liste: 
            
            datum_formatiert = datetime.datetime.strptime(datum_db, '%Y-%m-%d').strftime('%d.%m.%Y')
            datum_uhrzeit_text = f"{datum_formatiert}, von {uhrzeit_von} bis {uhrzeit_bis}" 
            
            summe_leistung = einzelbetrag 
            total_betrag += summe_leistung
            
            full_description = beschreibung
            if ' - ' in full_description:
                 full_description = full_description.split(' - ', 1)[-1] 
            
            leistung_block = template_text.replace('{{LEISTUNG_DATUM}}', datum_uhrzeit_text)
            leistung_block = leistung_block.replace('{{LEISTUNG_BESCHREIBUNG}}', full_description) 
            leistung_block = leistung_block.replace('{{LEISTUNG_SUMME}}', f"€ {summe_leistung:.2f}")
            
            # Trennlinie hinzufügen, um später den Abstand zu setzen
            leistung_block += '[[ENDE_BLOCK]]' + '\n\n' 
            
            gesamt_leistungs_text += leistung_block 
            
        if gesamt_leistungs_text:
            
            current_line_is_leistung = False
            lines_of_last_block = [] # Speichert die Zeilen des letzten vollständigen Blocks
            
            for zeile in gesamt_leistungs_text.split('\n'):
                content = zeile.strip() 
                
                if not content:
                    continue

                is_end_block_marker = '[[ENDE_BLOCK]]' in content
                
                # Wir sammeln die Zeilen des aktuellen Blocks, um später keep_with_next zu setzen
                if not is_end_block_marker:
                    lines_of_last_block.append(content)


                if content.startswith("Leistung:"):
                    current_line_is_leistung = True
                elif content.startswith("Datum:") or content.startswith("Betrag:"):
                    current_line_is_leistung = False
                
                
                # Fügen Sie den Paragraphen ein
                if content:
                    
                    # Wenn wir den Ende-Marker des Blocks sehen, ignorieren wir ihn für den Inhalt
                    if is_end_block_marker:
                        content = content.replace('[[ENDE_BLOCK]]', '').strip()
                        if not content:
                            continue

                    p_new = block_start_paragraph.insert_paragraph_before(content)
                    
                    # Allgemeine Formatierung
                    p_new.paragraph_format.space_before = Pt(0)
                    p_new.paragraph_format.space_after = Pt(0)
                    p_new.paragraph_format.line_spacing = 1.0
                    
                    for run in p_new.runs:
                        run.font.size = Pt(12)
                    
                    # Korrektur für Leistungs-Einzug (hängender Einzug)
                    if current_line_is_leistung:
                        p_new.paragraph_format.left_indent = LEISTUNG_INDENT_SIZE      
                        p_new.paragraph_format.first_line_indent = -LEISTUNG_INDENT_SIZE 

                    
                    # 💥 NEUE LOGIK FÜR SEITENUMBRUCH-SCHUTZ:
                    # Der Schutz muss für alle Zeilen eines Blocks aktiviert werden, AUSSER für die letzte Zeile.
                    if not is_end_block_marker:
                        # Setze Keep_with_next für alle Zeilen, die nicht die letzte Zeile des Blocks sind.
                        p_new.paragraph_format.keep_with_next = True
                    else:
                        # Für die letzte Zeile (mit dem ENDE_BLOCK Marker) setzen wir es auf False
                        p_new.paragraph_format.keep_with_next = False
                        
                        # Setze Abstand nach dem Leistungsblock (nach "Betrag:")
                        p_new.paragraph_format.space_after = SPACE_AFTER_LEISTUNG_BLOCK
                        
                        # Setze den Block zurück für den nächsten Durchgang
                        lines_of_last_block = []


            # Entferne die alten Platzhalter-Paragraphen
            block_start_paragraph.text = '' 
            for p in block_paragraphs:
                p._element.getparent().remove(p._element)
            block_end_paragraph.text = ''
        else:
             block_start_paragraph.text = '' 
             for p in block_paragraphs:
                 p._element.getparent().remove(p._element)
             block_end_paragraph.text = ''


    # Ersetzen des Gesamtbetrags
    total_betrag_str = f"{total_betrag:.2f}"
    
    for p in document.paragraphs:
        
        if '{{Gesamt_Betrag}}' in p.text:
            
            original_text = p.text
            p.text = '' 
            
            parts = original_text.split('{{Gesamt_Betrag}}')
            
            run_label = p.add_run(parts[0]) 
            run_label.bold = True
            run_label.font.size = Pt(12)
            
            run_value = p.add_run(total_betrag_str) 
            run_value.bold = True
            run_value.font.underline = WD_UNDERLINE.DOUBLE 
            run_value.font.size = Pt(12)
            
            if len(parts) > 1 and parts[1]:
                run_after = p.add_run(parts[1])
                run_after.bold = True 
                run_after.font.size = Pt(12)
        
        elif 'Gesamt:' in p.text:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(12)

    
    # --- 3. Letzte Zeilenabstands-Korrektur (Finaler Sweep) ---
    for p in document.paragraphs:
        if p is not HONORARNOTE_PARAGRAPH:
            p.paragraph_format.space_before = Pt(0)
            if p.paragraph_format.space_after is None: 
                p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

    # --- 3b. Überweisungstext Logik (NEU) ---
    # Wenn Überweisung = NEIN (0), ersetze den Bank-Block durch "Betrag dankend erhalten!"
    if not ueberweisung:
        for p in document.paragraphs:
            # Suche nach dem Standard-Text aus der Vorlage
            if "Bitte überweisen Sie den Betrag" in p.text:
                p.text = "" # Text löschen
                run = p.add_run("Betrag dankend erhalten!")
                run.bold = True
                run.font.size = Pt(12)
            # Falls IBAN/BIC in eigenen Paragraphen stehen oder Reste davon da sind
            if "BIC GIBA" in p.text or "IBAN AT73" in p.text:
                p.text = ""


    # --- 4. Leerzeilen an den gewünschten Stellen einfügen (Post-Processing) ---
    
    def insert_empty_line(target_p):
        """Fügt eine leere Zeile (Paragraph) VOR dem Ziel-Paragraph ein."""
        # Korrigierte Logik, da insert_paragraph_after nicht existiert
        try:
            p_new = target_p.insert_paragraph_before('')
            
            p_new.paragraph_format.space_before = Pt(0)
            p_new.paragraph_format.space_after = Pt(0)
            p_new.paragraph_format.line_spacing = 1.0
            p_new.add_run('').font.size = Pt(12) 
            
        except (ValueError, IndexError, AttributeError):
            pass

    # Wichtig: Die Leerzeilen werden jetzt VOR dem Paragraphen eingefügt.
    # Wenn Sie eine Leerzeile NACH dem Paragraphen wünschen, müssen Sie
    # den nächsten Paragraphen finden und DIESEM die Leerzeile VORANSTELLEN.
    
    # Hier verwenden wir die einfache Methode, die Leerzeile VOR dem Block einzufügen
    if DIAGNOSE_PLACEHOLDER_PARAGRAPH:
        insert_empty_line(DIAGNOSE_PLACEHOLDER_PARAGRAPH)
        
    if HONORARNOTE_PARAGRAPH:
        # Fügt eine Leerzeile VOR der Honorarnote ein (was nicht gewünscht ist).
        # Lassen Sie dies am besten weg, da die Abstände (space_after) am Anfang
        # des Dokuments besser gesteuert werden sollten.
        pass # insert_empty_line(HONORARNOTE_PARAGRAPH)
        
    if ORT_DATUM_PARAGRAPH:
        insert_empty_line(ORT_DATUM_PARAGRAPH)

    
    # 💥 --- 5. Final Diagnosis Format Override (DIE ULTIMATIVE KORREKTUR!) --- 💥
    if DIAGNOSE_PLACEHOLDER_PARAGRAPH:
        
        target_p = DIAGNOSE_PLACEHOLDER_PARAGRAPH
        p_element = target_p._element
        
        diagnosis_label = "Diagnose:   "
        diagnosis_value = replacements.get('{{Diagnose}}', diagnose)
        
        # 1. Neuen Paragraphen VOR dem alten Platzhalter einfügen
        p_new = target_p.insert_paragraph_before('')
        
        # 2. Formatierung für den Paragraphen festlegen (Abstände und neutraler Stil)
        p_new.paragraph_format.left_indent = Pt(0)      
        p_new.paragraph_format.first_line_indent = Pt(0) 
        p_new.paragraph_format.space_before = Pt(0)
        p_new.paragraph_format.space_after = Pt(0)
        p_new.paragraph_format.line_spacing = 1.0
        
        # 3. Text in zwei Runs einfügen und formatieren
        
        # Run für das Label ("Diagnose: ")
        run_label = p_new.add_run(diagnosis_label)
        run_label.bold = True
        run_label.font.size = Pt(12)
        
        # Run für den Wert (die Diagnose)
        run_value = p_new.add_run(diagnosis_value)
        run_value.bold = True 
        run_value.font.size = Pt(12) 
        
        # 4. Den alten Platzhalter-Paragraphen KOMPLETT aus dem Dokument entfernen
        p_element.getparent().remove(p_element)
        
        # Setze den Platzhalter auf den neuen Paragraphen, um die Leerzeile in Schritt 4 zu erhalten
        DIAGNOSE_PLACEHOLDER_PARAGRAPH = p_new
        
    # --- 6. Footer Zusammenhalten (Keep With Next) ---
    # Verhindert, dass der Block ab "Gesamt:" durch einen Seitenumbruch getrennt wird.
    footer_started = False
    for i, p in enumerate(document.paragraphs):
        if 'Gesamt:' in p.text:
            footer_started = True
        
        if footer_started:
            # keep_with_next für alle Zeilen bis zur vorletzten setzen
            if i < len(document.paragraphs) - 1:
                p.paragraph_format.keep_with_next = True


    # --- Speichern des Dokuments ---
    patient_folder_name = f"{nachname} {vorname}"
    
    output_folder = os.path.expanduser(CONFIG.get('PATIENT_BASE_DIR'))
    patient_output_path = os.path.join(output_folder, patient_folder_name)
    os.makedirs(patient_output_path, exist_ok=True)
    
    output_filename = f"Honorarnote Krankenkasse {template_data['BHAG_NUMMER']}.docx"
    output_path = os.path.join(patient_output_path, output_filename)

    document.save(output_path)
    return output_path